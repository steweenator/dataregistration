import os
import json
import shutil
import time
from datetime import datetime, timedelta
import sys
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import sqlite3
import configparser
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import re
import logging
import logging.handlers
from watchdog.observers import Observer
from watchdog.events import PatternMatchingEventHandler
import threading

# --- Global Import Check for pydicom and pynetdicom ---
PYNETDICOM_AVAILABLE = False
pydicom_module = None
Dataset = None
FileMetaDataset = None
AE = None
evt = None
AllStoragePresentationContexts = None
ALL_TRANSFER_SYNTAXES = None
ModalityWorklistInformationFind = None
Verification = None
StudyRootQueryRetrieveInformationModelFind = None
EncapsulatedPDFStorage = None
SecondaryCaptureImageStorage = None

try:
    import pydicom as pydicom_module_local
    from pydicom.dataset import Dataset as PydicomDatasetLocal, FileMetaDataset as PydicomFileMetaDatasetLocal
    from pynetdicom import (
        AE as PynetdicomAELocal, evt as pynetdicom_evt_local,
        AllStoragePresentationContexts as PynetdicomAllStoragePresentationContextsLocal,
        ALL_TRANSFER_SYNTAXES as PynetdicomALL_TRANSFER_SYNTAXESLocal
    )
    from pynetdicom.sop_class import (
        ModalityWorklistInformationFind as PynetdicomModalityWorklistInformationFindLocal,
        Verification as PynetdicomVerificationLocal,
        StudyRootQueryRetrieveInformationModelFind as SRQRFind,
        EncapsulatedPDFStorage as EncapsulatedPDFStorageLocal,
        SecondaryCaptureImageStorage as SecondaryCaptureImageStorageLocal
    )
    pydicom_module = pydicom_module_local
    Dataset = PydicomDatasetLocal
    FileMetaDataset = PydicomFileMetaDatasetLocal
    AE = PynetdicomAELocal
    evt = pynetdicom_evt_local
    AllStoragePresentationContexts = PynetdicomAllStoragePresentationContextsLocal
    ALL_TRANSFER_SYNTAXES = PynetdicomALL_TRANSFER_SYNTAXESLocal
    ModalityWorklistInformationFind = PynetdicomModalityWorklistInformationFindLocal
    Verification = PynetdicomVerificationLocal
    StudyRootQueryRetrieveInformationModelFind = SRQRFind
    EncapsulatedPDFStorage = EncapsulatedPDFStorageLocal
    SecondaryCaptureImageStorage = SecondaryCaptureImageStorageLocal
    PYNETDICOM_AVAILABLE = True
except ImportError:
    PYNETDICOM_AVAILABLE = False

# --- python-docx Import Check ---
DOCX_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("SUCCESS: python-docx imported successfully")
except ImportError:
    print("WARNING: python-docx not available. DOCX generation will be disabled.")
    Document = None

# --- Configuration and Constants ---
CONFIG_DIR = os.path.join(os.path.expanduser("~"), ".PatientRegistrationApp")
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.ini")
DB_FILE = os.path.join(CONFIG_DIR, "patient_data.db")
JOBS_FILE = os.path.join(CONFIG_DIR, "pending_dicom_jobs.json")
DEFAULT_DATA_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "PatientRegistrationData")
LOG_FILE = os.path.join(CONFIG_DIR, "app.log")

if not os.path.exists(CONFIG_DIR):
    os.makedirs(CONFIG_DIR, exist_ok=True)

# --- Setup Logging ---
# Main application log
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] (%(threadName)s) (%(module)s:%(lineno)d) %(message)s"
)
# Dedicated logger for the UI Activity Log
activity_logger = logging.getLogger('activity')
activity_logger.setLevel(logging.INFO)
# Prevent activity logs from propagating to the main file log
activity_logger.propagate = False

logging.info(f"--- Application Starting ---")
logging.info(f"Python Version: {sys.version.split()[0]}")
logging.info(f"PYNETDICOM_AVAILABLE: {PYNETDICOM_AVAILABLE}")
logging.info(f"DOCX_AVAILABLE: {DOCX_AVAILABLE}")

# --- Global Thread Lock for Job File ---
jobs_file_lock = threading.Lock()

MODALITIES = ["CT", "DX", "US", "MG", "MR", "Default"]
SCANNER_SUPPORTED_FILE_TYPES = [".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".bmp"]
EXTERNAL_REPORT_FILE_TYPES = (".pdf", ".doc", ".docx")
MODALITY_PATTERNS = ["*.pdf", "*.jpg", "*.jpeg", "*.png", "*.tiff", "*.bmp"]
EXTERNAL_REPORT_PATTERNS = ["*.pdf", "*.doc", "*.docx"]

# --- SQLite Database Helper Functions (Same as before) ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patient_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT NOT NULL,
            patient_id TEXT NOT NULL,
            accession_number TEXT NOT NULL UNIQUE,
            dob_yyyymmdd TEXT NOT NULL,
            sex TEXT NOT NULL,
            study_date TEXT NOT NULL,
            study_time TEXT NOT NULL,
            study_description TEXT NOT NULL,
            referred_from TEXT,
            modality TEXT NOT NULL,
            requesting_physician TEXT,
            requested_procedure_id TEXT,
            scheduled_station_ae_title TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_patient_id ON patient_records (patient_id)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_accession_number ON patient_records (accession_number)")
    conn.commit()
    conn.close()
    logging.info(f"Database initialized/checked at {DB_FILE}")

def db_execute(query, params=(), fetchone=False, fetchall=False, commit=False):
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute(query, params)
        if commit:
            conn.commit()
            return cursor.lastrowid if "INSERT" in query.upper() else True
        if fetchone:
            return cursor.fetchone()
        if fetchall:
            return cursor.fetchall()
        return True
    except sqlite3.Error as e:
        logging.error(f"Database error: {e} \nQuery: {query} \nParams: {params}")
        return None
    finally:
        conn.close()

# ... (all other db helper functions: add_patient_record_db, etc. are unchanged) ...
def add_patient_record_db(data_dict):
    query = '''
        INSERT INTO patient_records
        (patient_name, patient_id, accession_number, dob_yyyymmdd, sex,
         study_date, study_time, study_description, referred_from, modality,
         requesting_physician, requested_procedure_id, scheduled_station_ae_title)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    '''
    params = (
        data_dict.get("Patient Name"), data_dict.get("Patient ID"), data_dict.get("Accession Number"),
        data_dict.get("Date of Birth"), data_dict.get("Sex"), data_dict.get("Study Date"),
        data_dict.get("Study Time"), data_dict.get("Study Description"),
        data_dict.get("Referred From"), data_dict.get("Modality"),
        data_dict.get("Requesting Physician"), data_dict.get("Requested Procedure ID"),
        data_dict.get("Scheduled Station AE Title")
    )
    return db_execute(query, params, commit=True)

def update_patient_record_db(record_id, data_dict):
    """Update an existing patient record in the database"""
    query = '''
        UPDATE patient_records SET
        patient_name=?, patient_id=?, accession_number=?, dob_yyyymmdd=?, sex=?,
        study_date=?, study_time=?, study_description=?, referred_from=?, modality=?,
        requesting_physician=?, requested_procedure_id=?, scheduled_station_ae_title=?
        WHERE id=?
    '''
    params = (
        data_dict.get("Patient Name"), data_dict.get("Patient ID"), data_dict.get("Accession Number"),
        data_dict.get("Date of Birth"), data_dict.get("Sex"), data_dict.get("Study Date"),
        data_dict.get("Study Time"), data_dict.get("Study Description"),
        data_dict.get("Referred From"), data_dict.get("Modality"),
        data_dict.get("Requesting Physician"), data_dict.get("Requested Procedure ID"),
        data_dict.get("Scheduled Station AE Title"),
        record_id
    )
    return db_execute(query, params, commit=True)


def get_patient_by_id_db(patient_id_to_find):
    query = "SELECT * FROM patient_records WHERE patient_id = ? ORDER BY created_at DESC LIMIT 1"
    row = db_execute(query, (patient_id_to_find,), fetchone=True)
    return dict(row) if row else None

def get_patient_record_by_db_id(record_id):
    """Get a specific patient record by database ID"""
    query = "SELECT * FROM patient_records WHERE id = ?"
    row = db_execute(query, (record_id,), fetchone=True)
    return dict(row) if row else None

def check_duplicate_record_db(patient_name, patient_id, accession_number):
    now = datetime.now()
    threshold_dt = now - timedelta(hours=36)
    threshold_timestamp = threshold_dt.strftime("%Y-%m-%d %H:%M:%S")

    query_exact = """
        SELECT study_description, created_at FROM patient_records
        WHERE patient_id = ? AND accession_number = ? AND created_at > ?
        ORDER BY created_at DESC LIMIT 1
    """
    exact_match = db_execute(query_exact, (patient_id, accession_number, threshold_timestamp), fetchone=True)
    if exact_match:
        created_at_str = exact_match["created_at"]
        try:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S.%f")
        except ValueError:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S")
        diff = now - study_dt
        hrs = int(diff.total_seconds() // 3600)
        mins = int((diff.total_seconds() % 3600) // 60)
        return True, hrs, mins, study_dt.strftime("%b %d, %Y %H:%M:%S"), exact_match["study_description"], "Exact Patient ID and Accession match"

    query_general = """
        SELECT study_description, created_at FROM patient_records
        WHERE (patient_id = ? OR patient_name = ?) AND created_at > ?
        ORDER BY created_at DESC LIMIT 1
    """
    general_match = db_execute(query_general, (patient_id, patient_name, threshold_timestamp), fetchone=True)
    if general_match:
        created_at_str = general_match["created_at"]
        try:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S.%f")
        except ValueError:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S")
        diff = now - study_dt
        hrs = int(diff.total_seconds() // 3600)
        mins = int((diff.total_seconds() % 3600) // 60)
        return True, hrs, mins, study_dt.strftime("%b %d, %Y %H:%M:%S"), general_match["study_description"], "Patient Name or ID match"

    return False, None, None, None, None, None

def get_all_patient_records_db(search_term=""):
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    base_query_fields = "id, patient_name, patient_id, accession_number, dob_yyyymmdd, sex, study_date, study_time, study_description, referred_from, modality, requesting_physician, requested_procedure_id, scheduled_station_ae_title, created_at"
    if search_term:
        query = f"SELECT {base_query_fields} FROM patient_records WHERE patient_name LIKE ? OR patient_id LIKE ? OR accession_number LIKE ? OR study_description LIKE ? ORDER BY created_at DESC"
        like_term = f"%{search_term}%"
        cursor.execute(query, (like_term, like_term, like_term, like_term))
    else:
        query = f"SELECT {base_query_fields} FROM patient_records ORDER BY created_at DESC"
        cursor.execute(query)

    columns = [desc[0] for desc in cursor.description] if cursor.description else []
    all_data = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return all_data, columns

def get_distinct_values_for_combobox_db(field_name):
    rows = db_execute(f"SELECT DISTINCT {field_name} FROM patient_records WHERE {field_name} IS NOT NULL AND {field_name} != '' ORDER BY {field_name}", fetchall=True)
    return [row[0] for row in rows] if rows else []

# --- Job Queue Functions ---
def _read_jobs():
    with jobs_file_lock:
        if not os.path.exists(JOBS_FILE):
            return {}
        try:
            with open(JOBS_FILE, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {}

def _write_jobs(jobs_data):
    with jobs_file_lock:
        try:
            with open(JOBS_FILE, 'w') as f:
                json.dump(jobs_data, f, indent=4)
        except Exception as e:
            logging.error(f"Failed to write to jobs file {JOBS_FILE}: {e}")

# ... (MWLServerThread class is unchanged) ...
class MWLServerThread(threading.Thread):
    def __init__(self, app_config):
        super().__init__(daemon=True)
        self.app_config = app_config
        self.ae_instance = None
        self.server_running = False
        self.logger = logging.getLogger("mwl_scp")

    def handle_echo(self, pynetdicom_event):
        self.logger.info(f"C-ECHO request received from {pynetdicom_event.assoc.requestor.ae_title}@{pynetdicom_event.assoc.requestor.address}:{pynetdicom_event.assoc.requestor.port}")
        return 0x0000

    def handle_find(self, pynetdicom_event):
        self.logger.info(f"C-FIND request received from {pynetdicom_event.assoc.requestor.ae_title}@{pynetdicom_event.assoc.requestor.address}:{pynetdicom_event.assoc.requestor.port}")

        if not PYNETDICOM_AVAILABLE or Dataset is None or pydicom_module is None:
            self.logger.error("pydicom.Dataset or pydicom module not available, cannot process C-FIND.")
            yield 0xC001, None
            return

        req_identifier = pynetdicom_event.identifier
        self.logger.debug(f"C-FIND Request Identifier:\n{req_identifier}")

        sql_query = "SELECT * FROM patient_records WHERE 1=1"
        params = []

        if 'PatientName' in req_identifier and req_identifier.PatientName:
            pn = str(req_identifier.PatientName).replace('*', '%').replace('?', '_')
            sql_query += " AND patient_name LIKE ?"
            params.append(pn)
        if 'PatientID' in req_identifier and req_identifier.PatientID:
            sql_query += " AND patient_id = ?"
            params.append(str(req_identifier.PatientID))
        if 'AccessionNumber' in req_identifier and req_identifier.AccessionNumber:
            sql_query += " AND accession_number = ?"
            params.append(str(req_identifier.AccessionNumber))

        if 'ModalitiesInStudy' in req_identifier and req_identifier.ModalitiesInStudy :
            sql_query += " AND modality = ?"
            params.append(str(req_identifier.ModalitiesInStudy))
        elif hasattr(req_identifier, 'ScheduledProcedureStepSequence') and \
             req_identifier.ScheduledProcedureStepSequence and \
             'Modality' in req_identifier.ScheduledProcedureStepSequence[0] and \
             req_identifier.ScheduledProcedureStepSequence[0].Modality:
            sql_query += " AND modality = ?"
            params.append(str(req_identifier.ScheduledProcedureStepSequence[0].Modality))

        sps_start_date = ""
        if hasattr(req_identifier, 'ScheduledProcedureStepSequence') and \
           req_identifier.ScheduledProcedureStepSequence and \
           'ScheduledProcedureStepStartDate' in req_identifier.ScheduledProcedureStepSequence[0]:
            sps_start_date = req_identifier.ScheduledProcedureStepSequence[0].ScheduledProcedureStepStartDate

        if sps_start_date:
            if '-' in sps_start_date:
                start_date_range, end_date_range = sps_start_date.split('-')
                sql_query += " AND study_date BETWEEN ? AND ?"
                params.extend([start_date_range.strip(), end_date_range.strip()])
            else:
                sql_query += " AND study_date = ?"
                params.append(sps_start_date.strip())

        self.logger.debug(f"Executing SQL for C-FIND: {sql_query} with params: {params}")
        matching_records = db_execute(sql_query, tuple(params), fetchall=True)

        if matching_records is None:
            self.logger.error("Database error during C-FIND query.")
            yield 0xA700, None
            return

        self.logger.info(f"Found {len(matching_records)} records matching C-FIND criteria.")

        for record in matching_records:
            if not PYNETDICOM_AVAILABLE or Dataset is None or pydicom_module is None: break

            ds = Dataset()

            ds.PatientName = record["patient_name"] if record["patient_name"] else "UNKNOWN"
            ds.PatientID = record["patient_id"] if record["patient_id"] else "UNKNOWN"
            ds.PatientBirthDate = record["dob_yyyymmdd"] if record["dob_yyyymmdd"] else ""
            ds.PatientSex = record["sex"] if record["sex"] else ""

            ds.AccessionNumber = record["accession_number"] if record["accession_number"] else ""
            ds.ReferringPhysicianName = record["referred_from"] if record["referred_from"] else ""
            ds.StudyInstanceUID = pydicom_module.uid.generate_uid()

            ds.RequestingPhysician = record["requesting_physician"] if record["requesting_physician"] else ""
            ds.RequestedProcedureDescription = record["study_description"] if record["study_description"] else "UNKNOWN"
            ds.RequestedProcedureID = record["requested_procedure_id"] if record["requested_procedure_id"] else ""

            sps_item = Dataset()
            sps_item.ScheduledStationAETitle = record["scheduled_station_ae_title"] if record["scheduled_station_ae_title"] else (self.ae_instance.ae_title if self.ae_instance else "UNKNOWN_AE")
            sps_item.ScheduledProcedureStepStartDate = record["study_date"] if record["study_date"] else ""
            sps_item.ScheduledProcedureStepStartTime = record["study_time"] if record["study_time"] else ""
            sps_item.Modality = record["modality"] if record["modality"] else ""
            sps_item.ScheduledPerformingPhysicianName = ""
            sps_item.ScheduledProcedureStepDescription = record["study_description"] if record["study_description"] else "UNKNOWN"
            sps_item.ScheduledProcedureStepID = record["accession_number"]

            ds.ScheduledProcedureStepSequence = [sps_item]

            if hasattr(req_identifier, 'SpecificCharacterSet'):
                ds.SpecificCharacterSet = req_identifier.SpecificCharacterSet
            else:
                ds.SpecificCharacterSet = "ISO_IR 100"

            self.logger.debug(f"Yielding C-FIND response for Accession: {ds.AccessionNumber}")
            yield 0xFF00, ds

        self.logger.info("Finished processing C-FIND, yielding final success status.")
        yield 0x0000, None

    def _handle_acse_recv(self, event): # New handler for EVT_ACSE_RECV
        pdu_type_attr = None
        reason_diag_attr = 'N/A'
        result_attr = 'N/A'
        source_attr = 'N/A'
        if hasattr(event, 'pdu_type'):
            pdu_type_attr = event.pdu_type
            reason_diag_attr = event.reason_diag if hasattr(event, 'reason_diag') else 'N/A'
            result_attr = event.result if hasattr(event, 'result') else 'N/A'
            source_attr = event.source if hasattr(event, 'source') else 'N/A'
        elif hasattr(event, 'primitive') and hasattr(event.primitive, 'result_str'):
            pdu_type_attr = 0x03
            reason_diag_attr = event.primitive.information.get('Diagnostic', 'N/A') if hasattr(event.primitive, 'information') and isinstance(event.primitive.information, dict) else 'N/A'
            result_attr = event.primitive.result_str
            source_attr = event.primitive.source_str if hasattr(event.primitive, 'source_str') else 'N/A'
        elif hasattr(event, 'primitive') and hasattr(event.primitive, 'source_str'):
             pdu_type_attr = 0x07
             reason_diag_attr = event.primitive.information.get('Diagnostic', 'N/A') if hasattr(event.primitive, 'information') and isinstance(event.primitive.information, dict) else 'N/A'
             source_attr = event.primitive.source_str
        else:
            self.logger.info(f"ACSE PDU Received (structure unknown, cannot parse details): {event}")
            return

        ae_title_info = event.assoc.requestor.ae_title if event.assoc and event.assoc.requestor else 'Unknown AE'

        if pdu_type_attr == 0x03:
            self.logger.warning(f"Received A-ASSOCIATE-RJ (Rejection) from {ae_title_info}. Result: {result_attr}, Source: {source_attr}, Reason/Diag: {reason_diag_attr}")
        elif pdu_type_attr == 0x07:
             self.logger.error(f"Received A-ABORT PDU from {ae_title_info}. Source: {source_attr}, Reason/Diag: {reason_diag_attr}")
        else:
            self.logger.info(f"ACSE PDU Received (Type {pdu_type_attr if pdu_type_attr else 'Unknown'}): {event}")


    def run(self):
        if not PYNETDICOM_AVAILABLE or AE is None or ModalityWorklistInformationFind is None or Verification is None or ALL_TRANSFER_SYNTAXES is None or evt is None:
            self.logger.error("Cannot start MWL SCP server: Essential pynetdicom components not available globally.")
            self.server_running = False
            return

        ae_title_str = self.app_config.get("MWLServerConfig", "ae_title", fallback="PYMWLSCP")
        port = self.app_config.getint("MWLServerConfig", "port", fallback=11112)

        try:
            self.ae_instance = AE(ae_title=ae_title_str.encode('ascii'))
        except Exception as e_ae_init:
            self.logger.error(f"Failed to initialize AE for MWL Server: {e_ae_init}")
            self.server_running = False
            return

        self.ae_instance.add_supported_context(ModalityWorklistInformationFind, ALL_TRANSFER_SYNTAXES)
        self.ae_instance.add_supported_context(Verification, ALL_TRANSFER_SYNTAXES)

        handlers = [
            (evt.EVT_C_ECHO, self.handle_echo),
            (evt.EVT_C_FIND, self.handle_find),
            (evt.EVT_ACSE_RECV, self._handle_acse_recv),
            (evt.EVT_ABORTED, lambda event: self.logger.error(f"Connection Aborted: {event}")),
            (evt.EVT_REQUESTED, lambda event: self.logger.info(f"Association Requested: {event.assoc}")),
            (evt.EVT_ACCEPTED, lambda event: self.logger.info(f"Association Accepted by {event.assoc.acceptor.ae_title if event.assoc and event.assoc.acceptor else 'Unknown AE'}")),
            (evt.EVT_REJECTED, lambda event: self.logger.warning(f"Association Rejected by {event.assoc.acceptor.ae_title if event.assoc and event.assoc.acceptor else 'Unknown AE'}")),
            (evt.EVT_RELEASED, lambda event: self.logger.info(f"Association Released: {event.assoc}")),
        ]

        self.logger.info(f"Starting MWL SCP server on port {port} with AE Title {ae_title_str}...")
        self.server_running = True
        try:
            self.ae_instance.start_server(('', port), block=True, evt_handlers=handlers)
        except OSError as e:
            self.logger.error(f"OSError starting MWL SCP server (possibly port {port} in use): {e}")
            if tk._default_root and tk._default_root.winfo_exists():
                 messagebox.showerror("MWL Server Error", f"Could not start MWL server on port {port}.\nIs it already in use?\n\nError: {e}", parent=tk._default_root)
            else:
                 print(f"CRITICAL MWL Server Error (GUI not ready): Could not start MWL server on port {port}. Is it already in use? Error: {e}")
        except Exception as e:
            self.logger.exception(f"General exception in MWL SCP server: {e}")
        finally:
            self.server_running = False
            self.logger.info("MWL SCP server has stopped.")

    def stop_server(self):
        if self.ae_instance and self.server_running:
            self.logger.info("Attempting to shut down MWL SCP server...")
            try:
                self.ae_instance.shutdown()
            except Exception as e_shutdown:
                self.logger.error(f"Error during MWL server shutdown: {e_shutdown}")
            self.server_running = False
            self.logger.info("MWL SCP server shutdown initiated.")
        else:
            self.logger.info("MWL SCP server was not running or AE instance not initialized.")

# ... (WatchHandler class is unchanged) ...
class WatchHandler(PatternMatchingEventHandler):
    def __init__(self, patterns, callback, app_instance):
        super().__init__(patterns=patterns, ignore_directories=True, case_sensitive=False)
        self.callback = callback
        self.app = app_instance

    def on_created(self, event):
        path = event.src_path
        logging.info(f"Watchdog detected new file via on_created: {path}")
        try:
            if self.app and hasattr(self.app, 'root') and self.app.root.winfo_exists():
                 self.app.root.after(2000, lambda p=path: self._process_file(p))
            else:
                logging.warning(f"Root window not available/destroyed for delayed processing of {path}. Processing immediately.")
                self._process_file(path)
        except Exception as e:
            logging.exception(f"Error scheduling callback for {path}: {e}. Processing immediately.")
            self._process_file(path)

    def _ensure_local_temp_watch_dir(self):
        try:
            temp_watch_dir = os.path.join(CONFIG_DIR, "temp_watched_files")
            if not os.path.exists(temp_watch_dir):
                os.makedirs(temp_watch_dir, exist_ok=True)
                logging.info(f"_ensure_local_temp_watch_dir: Created temporary watch directory: {temp_watch_dir}")
            if os.access(temp_watch_dir, os.W_OK):
                return temp_watch_dir
            else:
                logging.error(f"_ensure_local_temp_watch_dir: Temporary watch directory is not writable: {temp_watch_dir}")
                return None
        except Exception as e:
            logging.error(f"_ensure_local_temp_watch_dir: Failed to create/access temporary watch directory: {e}")
            return None

    def _process_file(self, original_watched_path):
        max_copy_retries = 7
        retry_delay_seconds = 0.75

        logging.info(f"_process_file: Received original watched path {original_watched_path} for processing.")

        local_temp_dir = self._ensure_local_temp_watch_dir()
        if not local_temp_dir:
            logging.error(f"_process_file: Cannot proceed without a local temporary directory. Skipping {original_watched_path}.")
            return

        local_snapshot_path = None

        for i in range(max_copy_retries):
            if not os.path.exists(original_watched_path):
                logging.warning(f"_process_file: Attempt {i+1}/{max_copy_retries} - Original watched file {original_watched_path} does not exist. Retrying in {retry_delay_seconds}s.")
                time.sleep(retry_delay_seconds)
                continue

            try:
                original_size = os.path.getsize(original_watched_path)
                if original_size == 0:
                    logging.warning(f"_process_file: Attempt {i+1}/{max_copy_retries} - Original watched file {original_watched_path} is empty (size 0). Retrying in {retry_delay_seconds}s.")
                    time.sleep(retry_delay_seconds)
                    continue

                base_fname = os.path.basename(original_watched_path)
                timestamp_prefix = datetime.now().strftime('%Y%m%d%H%M%S%f')
                prospective_snapshot_fname = f"{timestamp_prefix}_{base_fname}"
                prospective_snapshot_path = os.path.join(local_temp_dir, prospective_snapshot_fname)

                logging.info(f"_process_file: Attempt {i+1}/{max_copy_retries} - Trying to copy (snapshot) {original_watched_path} (size {original_size} bytes) to local temp: {prospective_snapshot_path}")

                shutil.copy2(original_watched_path, prospective_snapshot_path)

                if os.path.exists(prospective_snapshot_path):
                    snapshot_size = os.path.getsize(prospective_snapshot_path)
                    if snapshot_size > 0:
                        if snapshot_size == original_size:
                            logging.info(f"_process_file: Successfully snapshotted {original_watched_path} to {prospective_snapshot_path}. Size matched ({snapshot_size} bytes).")
                            local_snapshot_path = prospective_snapshot_path
                            break
                        else:
                            logging.warning(f"_process_file: Snapshotted {original_watched_path} to {prospective_snapshot_path}, but size MISMATCH! Original: {original_size}, Snapshot: {snapshot_size}. Retrying.")
                            os.remove(prospective_snapshot_path)
                    else:
                         logging.warning(f"_process_file: Snapshot {prospective_snapshot_path} created but is 0-byte. Deleting snapshot and retrying.")
                         os.remove(prospective_snapshot_path)
                else:
                    logging.warning(f"_process_file: Snapshot copy to {prospective_snapshot_path} attempted, but destination file does not exist. Retrying.")

            except (FileNotFoundError, PermissionError, shutil.Error, OSError) as e:
                logging.error(f"_process_file: Attempt {i+1}/{max_copy_retries} - Error processing {original_watched_path}: {e}. Retrying.")
            except Exception as e:
                logging.error(f"_process_file: Attempt {i+1}/{max_copy_retries} - Generic error processing {original_watched_path}: {e}. Retrying.")

            if local_snapshot_path:
                break
            time.sleep(retry_delay_seconds)

        if local_snapshot_path:
            logging.info(f"_process_file: Original file {original_watched_path} successfully snapshotted. Proceeding with callback.")
            try:
                self.callback(local_snapshot_path)
            except Exception as e_callback:
                logging.exception(f"_process_file: Error in application callback for {local_snapshot_path} (original: {original_watched_path}): {e_callback}")
        else:
            logging.error(f"_process_file: Failed to create a valid local snapshot of {original_watched_path} after {max_copy_retries} retries. Skipping.")

# ... (Helper functions like get_script_directory, ensure_dir_exists, etc. are unchanged) ...
def get_script_directory():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def ensure_dir_exists(path_to_ensure):
    if path_to_ensure and not os.path.exists(path_to_ensure):
        try:
            os.makedirs(path_to_ensure, exist_ok=True)
            logging.info(f"Ensured directory exists: {path_to_ensure}")
        except Exception as e:
            logging.error(f"Failed to create directory {path_to_ensure}: {e}")

def normalize_path_for_config_section(path_str):
    if not path_str: return ""
    name = str(path_str).replace("\\", "_").replace("/", "_").replace(":", "_colon_").replace(" ", "_space_")
    name = re.sub(r'[^a-zA-Z0-9_.-]', '', name)
    return name

# ... (load_config, save_config, create_default_config are unchanged) ...
def load_config():
    config = configparser.ConfigParser(interpolation=None, allow_no_value=True)
    if not os.path.exists(CONFIG_FILE):
        logging.info(f"Config file not found at {CONFIG_FILE}. Creating default.")
        create_default_config(config)
    else:
        try:
            config.read(CONFIG_FILE)
            logging.info(f"Config file loaded from {CONFIG_FILE}")
        except (configparser.Error, Exception) as e_generic_read:
            logging.exception(f"Error reading config file {CONFIG_FILE}: {e_generic_read}. Will try to create default.")
            create_default_config(config)

    sections_to_check = [
        "Paths", "Preferences", "Paths.Output.DOCX.Modalities",
        "Paths.WatchFolders.Modalities", "UI.Labels", "SMTP",
        "Paths.ExternalReportWatchFolders", "EmailRecipients", "EmailTemplates",
        "MWLServerConfig", "DICOM.QueryPACS"
    ]
    for mod in MODALITIES:
        sections_to_check.append(f"DICOM.Destinations.{mod}")

    for section in sections_to_check:
        if not config.has_section(section):
            config.add_section(section)
            logging.info(f"Added missing section to config: {section}")

    default_db_path = DB_FILE
    default_template_path = os.path.join(DEFAULT_DATA_DIR, "Templates", "RADTEMPLATE.docx")
    default_general_docx_output = os.path.join(DEFAULT_DATA_DIR, "Reports", "_General")

    if not config.has_option("Paths", "db_file"): config.set("Paths", "db_file", default_db_path)
    if not config.has_option("Paths", "docx_template"): config.set("Paths", "docx_template", default_template_path)
    if not config.has_option("Paths", "general_docx_output_folder"): config.set("Paths", "general_docx_output_folder", default_general_docx_output)

    ensure_dir_exists(CONFIG_DIR)
    ensure_dir_exists(os.path.dirname(config.get("Paths", "db_file", fallback=default_db_path)))
    ensure_dir_exists(os.path.dirname(config.get("Paths", "docx_template", fallback=default_template_path)))
    ensure_dir_exists(config.get("Paths", "general_docx_output_folder", fallback=default_general_docx_output))

    return config

def save_config(config):
    ensure_dir_exists(CONFIG_DIR)
    try:
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
        logging.info(f"Configuration saved to {CONFIG_FILE}")
    except Exception as e:
        logging.exception(f"Error writing config file {CONFIG_FILE} during save_config")

def create_default_config(config):
    logging.info("Creating default configuration file.")
    config.clear()

    config['Paths'] = {
        'db_file': DB_FILE,
        'docx_template': os.path.join(DEFAULT_DATA_DIR, "Templates", "RADTEMPLATE.docx"),
        'general_docx_output_folder': os.path.join(DEFAULT_DATA_DIR, "Reports", "_General"),
        'general_watch_folder': "",
        'external_report_watch_folders_list': ""
    }
    config['Preferences'] = {
        'last_referred_from': '',
        'default_accession_prefix': 'CRH',
        'default_scheduled_station_ae': 'ANY_MODALITY',
        'color_theme': 'Default',
        'ui_style': 'System Default',
        'enable_tooltips': 'True',
        'ui_size': 'Default'
    }
    config['MWLServerConfig'] = {
        "enabled": "False",
        "ae_title": "PYREGMWL",
        "port": "11112"
    }
    config['DICOM.QueryPACS'] = {
        'ae_title': 'PACS_AE',
        'ip': '127.0.0.1',
        'port': '104'
    }
    config['Paths.Output.DOCX.Modalities'] = {}
    config['Paths.WatchFolders.Modalities'] = {}
    for mod in MODALITIES:
        config.add_section(f'DICOM.Destinations.{mod}')
        config.set(f"DICOM.Destinations.{mod}", "ae_title", f"{mod}_DEST_AE")
        config.set(f"DICOM.Destinations.{mod}", "ip", "127.0.0.1")
        config.set(f"DICOM.Destinations.{mod}", "port", "11104")
        config.set("Paths.Output.DOCX.Modalities", mod, os.path.join(DEFAULT_DATA_DIR, "Reports", mod))
        config.set("Paths.WatchFolders.Modalities", mod, os.path.join(DEFAULT_DATA_DIR, "WatchFolders", mod))

    default_labels = {
        "main_window_title": "Patient Registration & MWL Server", "patient_id": "Patient ID (e.g. 123456 AB):",
        "patient_name": "Patient Name:", "accession": "Accession (CRH[MODALITY]...):",
        "dob": "Date of Birth (DD/MM/YYYY or DDMMYYYY):", "sex": "Sex (M/F):",
        "study_description": "Study Description:", "referred_from": "Referred From:",
        "requesting_physician": "Requesting Physician:", "requested_procedure_id": "Requested Procedure ID:",
        "scheduled_station_ae": "Scheduled Station AE:",
        "attach_files_button": "Attach File(s)",
        "attachments_label_prefix": "Attachments:", "submit_button": "Register Patient", "clear_button": "Clear Form",
        "email_button": "Email Previous Report", "settings_window_title": "Settings",
        "view_data_window_title": "Patient Data Viewer",
        "view_served_worklist_title": "Served Worklist Viewer",
        "smtp_settings_tab_title": "SMTP Email",
        "smtp_server": "SMTP Server:", "smtp_port": "SMTP Port:", "smtp_user": "SMTP Username (optional):",
        "smtp_password": "SMTP Password (optional):", "smtp_sender_email": "Sender Email Address:",
        "smtp_use_tls": "Use TLS/STARTTLS Encryption", "smtp_test_button": "Test SMTP Settings",
        "email_picker_title": "Select Report to Email", "email_picker_button": "Compose Email for Selected Report",
        "email_composer_title": "Compose Email", "email_composer_to": "To:", "email_composer_subject": "Subject:",
        "email_composer_body": "Body:", "email_composer_attachments_label": "Attachments (auto-included):",
        "email_composer_send_button": "Send Email", "ext_reports_tab_title": "External Reports",
        "ext_reports_watch_folder_label": "Watch Folders for External Reports (PDF, DOC, DOCX):",
        "email_recipients_tab_title": "Email Recipients",
        "favorite_recipients_label": "Favorite Email Recipients:",
        "add_favorite_button": "Add Favorite",
        "remove_favorite_button": "Remove Selected Favorite",
        "max_recent_recipients_label": "Max Recent Recipients to Store:",
        "email_templates_tab_title": "Email Templates",
        "manage_email_templates_label": "Manage Email Templates:",
        "template_name_label": "Template Name:",
        "template_subject_label": "Subject Template:",
        "template_body_label": "Body Template:",
        "add_template_button": "Add New Template",
        "edit_template_button": "Edit Selected Template",
        "delete_template_button": "Delete Selected Template",
        "apply_template_button": "Apply Template",
        "available_placeholders_label": "Placeholders: {Patient Name}, {Modality}, {Study Description}, {Date}, {Report Filename}, {Attachment Count}, {All Attachment Names}",
        "email_composer_template_label": "Email Template:",
        "recent_recipients_combobox_label": "Recent/Favorites",
        "email_composer_select_recipients_label": "Select Recipients:",
        "email_composer_add_selected_button": "Add Selected to 'To:' Field",
        "ext_report_autosend_group_label": "Automatic Emailing for Selected Folder:",
        "ext_report_autosend_enable_label": "Enable Auto-Emailing for this Folder (PDF, DOC, DOCX):",
        "ext_report_autosend_recipients_label": "Auto-Send Recipient(s) (semicolon-separated):",
        "ext_report_autosend_add_favorite_button": "Add Favorite to Recipients",
        "ext_report_autosend_template_label": "Auto-Send Email Template:",
        "ext_report_autosend_custom_subject_label": "Custom Subject (if template is 'Custom'):",
        "ext_report_autosend_custom_body_label": "Custom Body (if template is 'Custom'):",
        "ext_report_autosend_placeholders_label": "Placeholders: {Filename}, {FolderPath}, {DateTime}",
        "appearance_tab_title": "Appearance Settings",
        "ui_style_engine_label": "UI Style Engine:",
        "color_palette_label": "Color Palette:",
        "app_mwl_server_tab_title": "This App as MWL Server",
        "app_mwl_server_enable_label": "Enable This Application as an MWL Server:",
        "app_mwl_server_ae_title_label": "This Application's AE Title:",
        "app_mwl_server_port_label": "Listening Port for DICOM Connections:"
    }
    config['UI.Labels'] = default_labels
    config['SMTP'] = {'server': '', 'port': '587', 'user': '', 'password': '', 'sender_email': '', 'use_tls': 'True'}
    config['Paths.ExternalReportWatchFolders'] = {}
    config['EmailRecipients'] = {'recent_list': '', 'favorite_list': '', 'max_recent': '10'}

    default_email_subject = "Report: {Patient Name} - {Modality} {Study Description} on {Date}"
    default_email_body = ("Dear Recipient,\n\nPlease find attached the report for patient: {Patient Name}.\n"
                          "Study: {Modality} - {Study Description}\nDate: {Date}\n\n"
                          "Report file: {Report Filename}\n"
                          "Number of attached files: {Attachment Count}\n"
                          "Attached files: {All Attachment Names}\n\n"
                          "Best regards,\n[Your Name/Clinic]")
    config['EmailTemplates'] = {
        'template_names': 'Default',
        'Default_subject': default_email_subject,
        'Default_body': default_email_body
    }

    ensure_dir_exists(CONFIG_DIR)
    ensure_dir_exists(DEFAULT_DATA_DIR)
    ensure_dir_exists(os.path.dirname(config.get("Paths", "db_file")))
    ensure_dir_exists(os.path.dirname(config.get("Paths", "docx_template")))
    ensure_dir_exists(config.get("Paths", "general_docx_output_folder"))

    general_watch_folder = config.get("Paths", "general_watch_folder", fallback="")
    if general_watch_folder: ensure_dir_exists(general_watch_folder)

    ext_report_folders_str = config.get("Paths", "external_report_watch_folders_list", fallback="")
    for folder in ext_report_folders_str.split(';'):
        if folder: ensure_dir_exists(os.path.normpath(os.path.expanduser(folder)))

    for mod in MODALITIES:
        ensure_dir_exists(config.get("Paths.Output.DOCX.Modalities", mod))
        mod_watch_folder = config.get("Paths.WatchFolders.Modalities", mod)
        if mod_watch_folder: ensure_dir_exists(mod_watch_folder)

    default_template_filename = "RADTEMPLATE.docx"
    source_template_path = os.path.join(get_script_directory(), default_template_filename)
    target_template_path = config.get("Paths", "docx_template")
    if os.path.exists(source_template_path) and not os.path.exists(target_template_path):
        try:
            ensure_dir_exists(os.path.dirname(target_template_path))
            shutil.copy2(source_template_path, target_template_path)
            logging.info(f"Copied default template '{default_template_filename}' to {target_template_path}")
        except Exception as e:
            logging.exception(f"Could not copy default template: {e}")
    elif not os.path.exists(source_template_path):
        logging.warning(f"Default template '{default_template_filename}' not found in script directory. User needs to provide one.")

    save_config(config)
    init_db()
    logging.info(f"Default configuration file created and saved at {CONFIG_FILE}. Default data folders in {DEFAULT_DATA_DIR}. Database at {DB_FILE}")

def get_modality_from_accession(accession_number):
    acc_upper = accession_number.upper()
    if acc_upper.startswith("CRHCT"): return "CT"
    if acc_upper.startswith("CRHDX"): return "DX"
    if acc_upper.startswith("CRHUS"): return "US"
    if acc_upper.startswith("CRHMG"): return "MG"
    if acc_upper.startswith("CRHMR"): return "MR"
    return None

def format_date_friendly(date_str_yyyymmdd):
    if not date_str_yyyymmdd: return "N/A"
    try:
        return datetime.strptime(date_str_yyyymmdd, "%Y%m%d").strftime("%b %d, %Y")
    except ValueError:
        return date_str_yyyymmdd

# --- UI Activity Log Handler ---
class TextHandler(logging.Handler):
    def __init__(self, text_widget):
        super().__init__()
        self.text_widget = text_widget
        self.setFormatter(logging.Formatter('%(asctime)s - %(message)s', '%H:%M:%S'))

    def emit(self, record):
        msg = self.format(record)
        def append():
            self.text_widget.configure(state='normal')
            self.text_widget.insert(tk.END, msg + '\n')
            self.text_widget.configure(state='disabled')
            self.text_widget.see(tk.END)
        # Safely update GUI from any thread
        self.text_widget.after(0, append)

# --- ToolTip Class (unchanged) ---
class ToolTip:
    def __init__(self, widget, text, app_config, main_app_instance):
        self.widget = widget
        self.text = text
        self.app_config = app_config
        self.main_app = main_app_instance
        self.tooltip = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event=None):
        if not self.app_config.getboolean("Preferences", "enable_tooltips", fallback=True):
            return
        if self.tooltip:
            self.tooltip.destroy()

        tooltip_bg = self.main_app.current_palette.get("tooltip_bg", "#FFFFE0")

        x_root, y_root = self.widget.winfo_rootx(), self.widget.winfo_rooty()
        y_final = y_root + self.widget.winfo_height() + 5

        self.tooltip = tk.Toplevel(self.widget)
        self.tooltip.wm_overrideredirect(True)

        label = ttk.Label(self.tooltip, text=self.text,
                          background=tooltip_bg,
                          relief="solid", borderwidth=1, padding=5,
                          wraplength=350)
        label.pack()

        self.tooltip.update_idletasks()
        tooltip_width = self.tooltip.winfo_width()

        final_x_pos = x_root + (self.widget.winfo_width() - tooltip_width) // 2

        screen_width = self.widget.winfo_screenwidth()
        if final_x_pos + tooltip_width > screen_width:
            final_x_pos = screen_width - tooltip_width - 10
        if final_x_pos < 5 :
            final_x_pos = 5

        self.tooltip.wm_geometry(f"+{int(final_x_pos)}+{int(y_final)}")

    def hide_tooltip(self, event=None):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None

# --- Main Application Class ---
class PatientRegistrationApp:
    def __init__(self, root_window):
        self.root = root_window
        init_db()
        self.config = load_config()
        self.manual_attachments = []
        self.current_palette = {}
        self.style = ttk.Style(self.root)

        self.manual_attach_label_var = tk.StringVar()
        self.status_var = tk.StringVar()

        self.main_frame = None
        self.status_bar = None
        self.mwl_server_thread = None
        self._observer = None

        self.apply_theme_and_styles()
        self.root.title(self.get_ui_label("main_window_title", "Patient Registration & MWL Server"))
        self.apply_ui_size()
        self.create_menu()

        # Initialize UI components
        self.create_widgets()
        self.create_status_bar()
        self.create_activity_log() # NEW

        self.load_combobox_values_from_db()
        self.entry_referred_from.set(self.config.get("Preferences", "last_referred_from", fallback=""))
        self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))
        self.entry_patient_id.focus()

        # Start background services
        self._start_realtime_watchers()
        self.update_email_button_state()
        self.start_mwl_server_if_configured()
        self._resume_pending_jobs() # NEW

        logging.info("PatientRegistrationApp initialized successfully.")

    def create_activity_log(self):
        log_frame = ttk.LabelFrame(self.root, text="Background Activity Log", padding=5)
        log_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=(0, 5))

        log_text = tk.Text(log_frame, height=8, state='disabled', wrap=tk.WORD,
                           bg=self.current_palette.get("entry_bg"),
                           fg=self.current_palette.get("entry_fg"))
        log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=log_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        log_text['yscrollcommand'] = scrollbar.set

        # Configure the dedicated logger to use the TextHandler
        text_handler = TextHandler(log_text)
        activity_logger.addHandler(text_handler)

    # ... (start/stop/check MWL server methods are unchanged) ...
    def start_mwl_server_if_configured(self):
        if not PYNETDICOM_AVAILABLE:
            self.update_status("MWL Server disabled: pynetdicom/pydicom components not found.", True, 0)
            return

        if self.config.getboolean("MWLServerConfig", "enabled", fallback=False):
            if self.mwl_server_thread and self.mwl_server_thread.is_alive():
                logging.info("MWL Server thread already running.")
                return

            self.mwl_server_thread = MWLServerThread(self.config)
            self.mwl_server_thread.start()
            self.root.after(1000, self.check_mwl_server_status)
        else:
            logging.info("MWL Server is disabled in configuration.")
            self.update_status("MWL Server is disabled.", False, 5000)

    def check_mwl_server_status(self):
        if not PYNETDICOM_AVAILABLE:
            self.update_status("MWL Server cannot run: pynetdicom/pydicom missing.", True, 0)
            return

        if self.mwl_server_thread and self.mwl_server_thread.is_alive() and self.mwl_server_thread.server_running:
            ae_title = self.config.get("MWLServerConfig", "ae_title", fallback="N/A")
            port = self.config.get("MWLServerConfig", "port", fallback="N/A")
            self.update_status(f"MWL Server running: {ae_title} on port {port}", False, 0)
            logging.info(f"MWL Server confirmed running: {ae_title} on port {port}")
        elif self.config.getboolean("MWLServerConfig", "enabled", fallback=False):
            self.update_status("MWL Server failed to start or stopped. Check logs.", True, 0)
            logging.error("MWL Server was enabled but is not running. Check logs for errors (e.g., port conflict or pynetdicom issue).")

    def stop_mwl_server(self):
        if self.mwl_server_thread and self.mwl_server_thread.is_alive():
            logging.info("Stopping MWL Server thread...")
            self.mwl_server_thread.stop_server()
            self.mwl_server_thread.join(timeout=5)
            if self.mwl_server_thread.is_alive():
                logging.warning("MWL Server thread did not stop gracefully after 5 seconds.")
            else:
                logging.info("MWL Server thread stopped.")
        self.mwl_server_thread = None
        self.update_status("MWL Server stopped.", False, 5000)

    def _start_realtime_watchers(self):
        if self._observer is not None:
            try:
                self._observer.stop()
                self._observer.join(timeout=1)
            except Exception as e:
                logging.warning(f"Could not stop previous Watchdog observer: {e}")

        self._observer = Observer()
        watched_folders = set()
        all_watch_paths = self.get_all_watch_folders_with_owners()

        # RECURSIVE WATCHERS
        for path, owner_modality in all_watch_paths.items():
            if path and os.path.isdir(path) and path not in watched_folders:
                try:
                    # Differentiate callbacks based on owner
                    if owner_modality == "ExternalReport":
                        handler = WatchHandler(EXTERNAL_REPORT_PATTERNS, self._process_auto_send_for_external_report_file, self)
                        log_msg = f"Scheduled recursive watchdog for external report folder: {path}"
                    else: # Modality or General
                        handler = WatchHandler(MODALITY_PATTERNS, lambda p, m=owner_modality: self._on_new_modality_file(p, m), self)
                        log_msg = f"Scheduled recursive watchdog for modality folder ({owner_modality or 'General'}): {path}"

                    self._observer.schedule(handler, path, recursive=True)
                    watched_folders.add(path)
                    logging.info(log_msg)
                except Exception as e:
                    logging.error(f"Failed to schedule watchdog for {path}: {e}")
            elif path and not os.path.isdir(path):
                logging.warning(f"Configured watch folder is not a valid directory: {path}")

        if not self._observer.emitters:
            logging.warning("Watchdog observer has no paths to watch.")
        else:
            try:
                self._observer.start()
                logging.info("Watchdog observer started successfully.")
            except Exception as e:
                logging.exception("Failed to start Watchdog observer.")

    def get_all_watch_folders_with_owners(self):
        """Returns a dict of all watch paths to their owner (modality code or 'ExternalReport')."""
        paths = {}
        for mod in MODALITIES:
            path = self.get_modality_specific_path("Paths.WatchFolders.Modalities", mod)
            if path: paths[path] = mod

        gen_path = self.config.get("Paths", "general_watch_folder", fallback="")
        if gen_path: paths[gen_path] = None # None represents 'General'

        ext_paths = self.get_all_external_report_watch_folders()
        for path in ext_paths:
            paths[path] = "ExternalReport"
        return paths

    # ... (_on_new_modality_file is unchanged) ...
    def _on_new_modality_file(self, local_snapshot_filepath, modality_code):
        if not os.path.exists(local_snapshot_filepath) or os.path.getsize(local_snapshot_filepath) == 0:
            logging.error(f"_on_new_modality_file: Received empty/missing snapshot path: {local_snapshot_filepath}. Skipping.")
            return

        snapshot_basename = os.path.basename(local_snapshot_filepath)
        try:
            original_fname_for_dest = "_".join(snapshot_basename.split("_")[1:])
            if not original_fname_for_dest: original_fname_for_dest = snapshot_basename
        except IndexError:
             original_fname_for_dest = snapshot_basename

        logging.info(f"Processing snapshot: {local_snapshot_filepath} (modality={modality_code or 'General'}).")

        if local_snapshot_filepath not in self.manual_attachments:
            self.manual_attachments.append(local_snapshot_filepath)
            self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")

    # ... (shutdown method is unchanged) ...
    def shutdown(self):
        logging.info("Application shutdown sequence initiated.")
        self.stop_mwl_server()

        if self._observer and self._observer.is_alive():
            try:
                self._observer.stop()
                self._observer.join(timeout=2)
                logging.info("Watchdog observer stopped.")
            except Exception as e:
                logging.exception("Exception during Watchdog observer shutdown.")
        else:
            logging.info("Watchdog observer was not running or already stopped.")

        logging.info("Application shutdown complete.")
        
    # ... (get_ui_label, apply_ui_size, apply_theme_and_styles, create_menu, etc., are unchanged) ...
    def get_ui_label(self, key, default_text=""):
        return self.config.get("UI.Labels", key, fallback=default_text)

    def apply_ui_size(self):
        size_setting = self.config.get("Preferences", "ui_size", fallback="Default")
        sizes = {
            "Very Compact": "700x580",
            "Compact": "700x610",
            "Default": "750x780", # Increased height for activity log
            "Large": "850x850",
            "Extra Large": "950x900"
        }
        self.root.geometry(sizes.get(size_setting, sizes["Default"]))

    def apply_theme_and_styles(self):
        selected_theme_name = self.config.get("Preferences", "color_theme", fallback="Default")
        selected_ui_style = self.config.get("Preferences", "ui_style", fallback="System Default")
        logging.info(f"Applying UI Style: {selected_ui_style}, Color Theme: {selected_theme_name}")

        base_ttk_theme = "clam"
        if selected_ui_style == "Clam (Modern)": base_ttk_theme = "clam"
        elif selected_ui_style == "Alt (Modern-ish)": base_ttk_theme = "alt"
        elif selected_ui_style == "Default (Classic)": base_ttk_theme = "default"
        elif selected_ui_style == "Classic (Older)": base_ttk_theme = "classic"
        elif selected_ui_style == "System Default":
            available_themes = self.style.theme_names()
            if 'vista' in available_themes and sys.platform == "win32": base_ttk_theme = 'vista'
            elif 'aqua' in available_themes and sys.platform == "darwin": base_ttk_theme = 'aqua'
            elif 'clam' in available_themes: base_ttk_theme = 'clam'
            else: base_ttk_theme = self.style.theme_use()

        try:
            self.style.theme_use(base_ttk_theme)
        except tk.TclError:
            logging.warning(f"TTK theme '{base_ttk_theme}' not found, falling back to 'clam'.")
            base_ttk_theme = "clam"
            self.style.theme_use(base_ttk_theme)

        themes = {
            "Default": {"bg": "#F0F0F0", "fg": "black", "entry_bg": "white", "entry_fg": "black", "button_bg": "#E0E0E0", "button_fg": "black", "label_fg": "black", "frame_bg": "#F0F0F0", "header_fg": "#0078D7", "tooltip_bg": "#FFFFE0"},
            "Light Blue": {"bg": "#E6F3FF", "fg": "#003366", "entry_bg": "#FFFFFF", "entry_fg": "#003366", "button_bg": "#B3D9FF", "button_fg": "#003366", "label_fg": "#004C99", "frame_bg": "#E6F3FF", "header_fg": "#0052A3", "tooltip_bg": "#D9EDFF"},
            "Dark": {"bg": "#2E2E2E", "fg": "#E0E0E0", "entry_bg": "#3C3C3C", "entry_fg": "#E0E0E0", "button_bg": "#505050", "button_fg": "#FFFFFF", "label_fg": "#C0C0C0", "frame_bg": "#2E2E2E", "header_fg": "#58A6FF", "tooltip_bg": "#4A4A4A"},
            "High Contrast": {"bg": "white", "fg": "black", "entry_bg": "white", "entry_fg": "black", "button_bg": "black", "button_fg": "white", "label_fg": "black", "frame_bg": "white", "header_fg": "blue", "tooltip_bg": "#FFFFE0"},
            "Mint Green": {"bg": "#E0F2F1", "fg": "#004D40", "entry_bg": "#FFFFFF", "entry_fg": "#004D40", "button_bg": "#A7FFEB", "button_fg": "#004D40", "label_fg": "#00695C", "frame_bg": "#E0F2F1", "header_fg": "#00796B", "tooltip_bg": "#B2DFDB"},
            "Lavender": {"bg": "#F3E5F5", "fg": "#4A148C", "entry_bg": "#FFFFFF", "entry_fg": "#4A148C", "button_bg": "#E1BEE7", "button_fg": "#4A148C", "label_fg": "#6A1B9A", "frame_bg": "#F3E5F5", "header_fg": "#8E24AA", "tooltip_bg": "#E1BEE7"}
        }
        palette = themes.get(selected_theme_name, themes["Default"])
        self.current_palette = palette

        self.root.configure(bg=palette["bg"])
        self.style.configure('.', background=palette["bg"], foreground=palette["fg"])
        self.style.configure('TLabel', font=('Helvetica', 11), padding=5, background=palette["bg"], foreground=palette["label_fg"])
        self.style.configure('TButton', font=('Helvetica', 11, 'bold'), padding=5, background=palette["button_bg"], foreground=palette["button_fg"])
        self.style.map('TButton', background=[('active', palette.get("button_active_bg", palette["button_bg"]))])
        self.style.configure('TEntry', font=('Helvetica', 11), padding=5)
        self.style.map('TEntry', fieldbackground=[('!focus', palette["entry_bg"]), ('focus', palette["entry_bg"])], foreground=[('!focus', palette["entry_fg"]), ('focus', palette["entry_fg"])])
        self.style.configure('TCombobox', font=('Helvetica', 11), padding=5)
        self.style.map('TCombobox', fieldbackground=[('readonly', palette["entry_bg"]), ('!readonly', palette["entry_bg"])], foreground=[('readonly', palette["entry_fg"]), ('!readonly', palette["entry_fg"])])
        self.root.option_add('*TCombobox*Listbox.background', palette["entry_bg"])
        self.root.option_add('*TCombobox*Listbox.foreground', palette["entry_fg"])
        self.root.option_add('*TCombobox*Listbox.selectBackground', palette.get("header_fg", "#0078D7"))
        self.root.option_add('*TCombobox*Listbox.selectForeground', palette.get("button_fg", "white"))
        self.style.configure('Header.TLabel', font=('Helvetica', 14, 'bold'), foreground=palette["header_fg"], background=palette["bg"])
        self.style.configure('Custom.TFrame', background=palette["frame_bg"])
        self.style.configure('Status.TLabel', background=palette.get("status_bg", palette["bg"]), foreground=palette["fg"])
        self.style.configure('Treeview.Heading', font=('Helvetica', 10, 'bold'), background=palette.get("button_bg"), foreground=palette.get("button_fg"))
        self.style.map("Treeview.Heading", background=[('active', palette.get("header_fg"))])
        self.style.configure("Treeview", rowheight=25, font=('Helvetica', 10),
                             background=palette.get("entry_bg"), foreground=palette.get("entry_fg"),
                             fieldbackground=palette.get("entry_bg"))
        self.style.map("Treeview",
                       background=[('selected', palette.get("header_fg", "#0078D7"))],
                       foreground=[('selected', palette.get("button_fg", "white"))])
        self.root.option_add('*Listbox.background', palette["entry_bg"])
        self.root.option_add('*Listbox.foreground', palette["entry_fg"])
        self.root.option_add('*Listbox.selectBackground', palette.get("header_fg", "#0078D7"))
        self.root.option_add('*Listbox.selectForeground', palette.get("button_fg", "white"))
        self.root.option_add('*Text.background', palette["entry_bg"])
        self.root.option_add('*Text.foreground', palette["entry_fg"])
        self.root.option_add('*Text.insertBackground', palette["entry_fg"])
        self.root.option_add('*Text.selectBackground', palette.get("header_fg", "#0078D7"))
        self.root.option_add('*Text.selectForeground', palette.get("button_fg", "white"))

        if self.main_frame:
            self.main_frame.destroy()
            self.create_widgets()
        if hasattr(self, 'status_bar') and self.status_bar:
            self.status_bar.destroy()
            self.create_status_bar()

    def create_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        menu_bg = self.current_palette.get("frame_bg", self.root.cget('bg'))
        menu_fg = self.current_palette.get("fg", "black")
        file_menu = tk.Menu(menubar, tearoff=0, bg=menu_bg, fg=menu_fg)
        file_menu.add_command(label="Settings", command=self.open_settings_window)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        menubar.add_cascade(label="File", menu=file_menu)
        view_menu = tk.Menu(menubar, tearoff=0, bg=menu_bg, fg=menu_fg)
        view_menu.add_command(label=self.get_ui_label("view_data_window_title", "View Patient Data"),
                              command=self.view_patient_data_window)
        view_menu.add_command(label=self.get_ui_label("view_served_worklist_title", "View Served Worklist"),
                              command=self.open_served_worklist_viewer)
        menubar.add_cascade(label="View", menu=view_menu)

    def create_widgets(self):
        if self.main_frame and self.main_frame.winfo_exists():
            self.main_frame.destroy()

        self.main_frame = ttk.Frame(self.root, padding="15", style='Custom.TFrame')
        self.main_frame.pack(side=tk.TOP, expand=True, fill=tk.BOTH)
        self.main_frame.columnconfigure(1, weight=1)

        self.create_fields()

        attach_frame_row = self.num_fields
        attach_frame = ttk.Frame(self.main_frame, style='Custom.TFrame')
        attach_frame.grid(row=attach_frame_row, column=0, columnspan=2, pady=8, sticky=tk.EW)

        self.attach_button = ttk.Button(attach_frame,
                                        text=self.get_ui_label("attach_files_button", "Attach File(s)"),
                                        command=self.select_manual_attachments, width=18)
        self.attach_button.pack(side=tk.LEFT, padx=(0, 5))
        ToolTip(self.attach_button, "Manually select files to be included with the generated DOCX report.", self.config, self)

        self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")
        manual_attach_display_label = ttk.Label(attach_frame, textvariable=self.manual_attach_label_var)
        manual_attach_display_label.pack(side=tk.LEFT, padx=(10, 0))

        self.create_buttons()

    # ... (create_fields, create_buttons, select_manual_attachments are unchanged) ...
    def create_fields(self):
        fields_config = [
            ("patient_id", "<FocusOut>", self.on_patient_id_change, "Unique patient ID. Auto-formats. Fetches existing data on FocusOut."),
            ("patient_name", None, None, "Full name of the patient."),
            ("accession", None, None, "Accession Number (e.g., CRHCT123). Must be unique."),
            ("dob", None, None, "Patient's date of birth (DD/MM/YYYY or DDMMYYYY)."),
            ("sex", None, None, "Sex (M for Male, F for Female)."),
            ("study_description", None, None, "Description of the study or examination."),
            ("referred_from", None, None, "Referring clinic or doctor."),
            ("requesting_physician", None, None, "Name of the physician requesting the study.")
        ]
        self.field_widgets = {}
        for i, (key, bind_event, bind_function, tooltip_text) in enumerate(fields_config):
            label_text = self.get_ui_label(key, key.replace("_", " ").title() + ":")
            label = ttk.Label(self.main_frame, text=label_text)
            label.grid(row=i, column=0, padx=5, pady=7, sticky=tk.W)
            entry_var_name = f"entry_{key}"
            if key in ["study_description", "referred_from", "requesting_physician"]:
                entry_widget = ttk.Combobox(self.main_frame, width=38, font=('Helvetica', 11))
            else:
                entry_widget = ttk.Entry(self.main_frame, width=40, font=('Helvetica', 11))
            entry_widget.grid(row=i, column=1, padx=5, pady=7, sticky=tk.EW)
            setattr(self, entry_var_name, entry_widget)
            self.field_widgets[entry_var_name] = entry_widget
            ToolTip(entry_widget, tooltip_text, self.config, self)
            if bind_event and bind_function:
                entry_widget.bind(bind_event, bind_function)
        self.num_fields = len(fields_config)

    def create_buttons(self):
        button_frame_row = self.num_fields + 1
        button_frame = ttk.Frame(self.main_frame, style='Custom.TFrame')
        button_frame.grid(row=button_frame_row, column=0, columnspan=2, pady=12, sticky=tk.EW)
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)
        button_frame.columnconfigure(2, weight=1)
        self.submit_button = ttk.Button(button_frame,
                                        text=self.get_ui_label("submit_button", "Register Patient"),
                                        command=self.submit_form, width=25)
        self.submit_button.grid(row=0, column=0, padx=5, pady=5, sticky=tk.E)
        ToolTip(self.submit_button, "Register the patient and start background DICOM workflow if applicable.", self.config, self)
        self.clear_button = ttk.Button(button_frame,
                                       text=self.get_ui_label("clear_button", "Clear Form"),
                                       command=self.confirm_clear_form, width=15)
        self.clear_button.grid(row=0, column=1, padx=5, pady=5)
        ToolTip(self.clear_button, "Clear all input fields and staged attachments.", self.config, self)
        self.email_button = ttk.Button(button_frame,
                                       text=self.get_ui_label("email_button", "Email Previous Report"),
                                       command=self.trigger_email_report_picker, width=20)
        self.email_button.grid(row=0, column=2, padx=5, pady=5, sticky=tk.W)
        ToolTip(self.email_button, "Select a previously generated report to email.", self.config, self)

    def select_manual_attachments(self):
        file_types = (
            ("All files", "*.*"),
            ("PDF files", "*.pdf"),
            ("Image files", "*.jpg;*.jpeg;*.png;*.tiff;*.bmp"),
            ("DICOM files", "*.dcm"),
            ("Document files", "*.doc;*.docx;*.txt")
        )
        filepaths = filedialog.askopenfilenames(
            title="Select Files to Attach",
            parent=self.root,
            filetypes=file_types
        )
        if filepaths:
            newly_added_count = 0
            for fp in filepaths:
                if fp not in self.manual_attachments:
                    self.manual_attachments.append(fp)
                    newly_added_count +=1
            if newly_added_count > 0:
                self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")
                self.update_status(f"{newly_added_count} file(s) manually staged for attachment.")

    def create_status_bar(self):
        if hasattr(self, 'status_bar') and self.status_bar and self.status_bar.winfo_exists():
            self.status_bar.destroy()
        self.status_var.set("Ready")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var,
                                    relief=tk.SUNKEN, anchor=tk.W, padding=(5,2),
                                    style="Status.TLabel")
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def update_status(self, message, is_error=False, duration=5000):
        if not (hasattr(self, 'status_var') and self.status_var):
            return
        self.status_var.set(message)
        if is_error:
            logging.error(f"Status Update (Error): {message}")
        else:
            logging.info(f"Status Update: {message}")
        if duration > 0 and '...' not in message:
            self.root.after(duration, lambda: self.status_var.set("Ready") if self.status_var.get() == message else None)

    # ... (get_modality_specific_path, load_combobox_values_from_db, check_for_duplicates, generate_docx_report, process_attachments are unchanged) ...
    def get_modality_specific_path(self, base_section_key, modality_code):
        if modality_code and self.config.has_option(base_section_key, modality_code):
            path = self.config.get(base_section_key, modality_code, fallback=None)
            if path: return os.path.normpath(os.path.expanduser(path))
        if self.config.has_option(base_section_key, "Default"):
            path = self.config.get(base_section_key, "Default", fallback=None)
            if path: return os.path.normpath(os.path.expanduser(path))
        if "Output.DOCX" in base_section_key:
            return os.path.normpath(os.path.expanduser(self.config.get("Paths", "general_docx_output_folder", fallback="")))
        elif "WatchFolders" in base_section_key:
             path = self.config.get("Paths", "general_watch_folder", fallback="")
             return os.path.normpath(os.path.expanduser(path)) if path else ""
        return ""

    def load_combobox_values_from_db(self):
        try:
            referred_from_values = get_distinct_values_for_combobox_db("referred_from")
            study_desc_values = get_distinct_values_for_combobox_db("study_description")
            req_phys_values = get_distinct_values_for_combobox_db("requesting_physician")
            sched_ae_values = get_distinct_values_for_combobox_db("scheduled_station_ae_title")

            if hasattr(self, 'entry_referred_from'):
                self.entry_referred_from['values'] = sorted(list(referred_from_values))
            if hasattr(self, 'entry_study_description'):
                self.entry_study_description['values'] = sorted(list(study_desc_values))
            if hasattr(self, 'entry_requesting_physician'):
                self.entry_requesting_physician['values'] = sorted(list(req_phys_values))
        except Exception as e:
            self.update_status(f"Error loading combobox values from DB: {e}", True)

    def check_for_duplicates(self, patient_name, patient_id, accession_number):
        is_dup, hrs, mins, prev_dt, prev_desc, match_type = check_duplicate_record_db(patient_name, patient_id, accession_number)
        if is_dup:
            logging.info(f"Duplicate check: Found recent entry for {patient_name}/{patient_id} (Type: {match_type})")
        return is_dup, hrs, mins, prev_dt, prev_desc

    def generate_docx_report(self, data_dict, modality_code, patient_specific_base_path):
        if not DOCX_AVAILABLE: return None
        template_path = self.config.get("Paths", "docx_template", fallback="")
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("DOCX Template Error", f"DOCX template not found at:\n{template_path}", parent=self.root)
            return None
        try:
            doc = Document(template_path)
        except Exception as e:
            messagebox.showerror("DOCX Load Error", f"Error loading template '{os.path.basename(template_path)}':\n{e}", parent=self.root)
            return None

        dob_f = format_date_friendly(data_dict.get("Date of Birth", ""))
        study_date_f = format_date_friendly(data_dict.get("Study Date", ""))
        sex_val = data_dict.get("Sex", "")
        sex_f = "Male" if sex_val == "M" else "Female" if sex_val == "F" else sex_val

        replacements = {
            '{Patient Name}': data_dict.get("Patient Name", ""), '{Docket Number}': data_dict.get("Patient ID", ""),
            '{Date of Birth}': dob_f, '{Accession Number}': data_dict.get("Accession Number", ""),
            '{Study Description}': data_dict.get("Study Description", ""), '{Referring Physician}': data_dict.get("Referred From", ""),
            '{Requesting Physician}': data_dict.get("Requesting Physician", ""), '{Study Date}': study_date_f,
            '{Date of Exam}': study_date_f, '{Clinic Referred From}': data_dict.get("Referred From", ""),
            '{Body Part Done}': data_dict.get("Study Description", ""), '{Modality}': data_dict.get("Modality", modality_code),
            '{Modality Done}': data_dict.get("Modality", modality_code), '{Sex}': sex_f
        }

        for p in doc.paragraphs:
            for r in p.runs:
                for k, v in replacements.items():
                    if k in r.text: r.text = r.text.replace(k, str(v))
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p_cell in cell.paragraphs:
                        for r_cell in p_cell.runs:
                            for k,v in replacements.items():
                                if k in r_cell.text: r_cell.text = r_cell.text.replace(k,str(v))

        safe_name = "".join(c if c.isalnum() else "_" for c in data_dict.get("Patient Name", "UnknownPatient"))
        safe_desc = "".join(c if c.isalnum() else "_" for c in data_dict.get("Study Description", "NoDesc"))[:30]
        study_date_filename_part = study_date_f.replace(',', '').replace(' ', '_') if study_date_f != "N/A" else datetime.now().strftime("%b_%d_%Y")
        fname = f"{safe_name}_{modality_code}_{safe_desc}_{study_date_filename_part}_REPORT.docx"
        output_file_path = os.path.join(patient_specific_base_path, fname)

        try:
            doc.save(output_file_path)
            self.update_status(f"DOCX report generated: {fname}")
            return output_file_path
        except Exception as e:
            messagebox.showerror("DOCX Save Error", f"Error saving DOCX report '{fname}':\n{e}", parent=self.root)
            return None

    def process_attachments(self, patient_id_for_match, modality_code, patient_specific_report_folder):
        all_copied_to_report_folder = []
        copied_count_manual = 0
        if self.manual_attachments:
            ensure_dir_exists(patient_specific_report_folder)
            attachments_to_process = list(self.manual_attachments)
            processed_indices = []
            for idx, src_path in enumerate(attachments_to_process):
                is_local_snapshot = CONFIG_DIR in src_path and "temp_watched_files" in src_path and os.path.dirname(src_path).endswith("temp_watched_files")
                if os.path.isfile(src_path):
                    try:
                        if is_local_snapshot:
                            snapshot_basename = os.path.basename(src_path)
                            try:
                                final_filename = "_".join(snapshot_basename.split("_")[1:])
                                if not final_filename: final_filename = snapshot_basename
                            except IndexError:
                                final_filename = snapshot_basename
                        else:
                            final_filename = os.path.basename(src_path)
                        dest_path = os.path.join(patient_specific_report_folder, final_filename)
                        if os.path.exists(dest_path) and os.path.abspath(src_path) != os.path.abspath(dest_path):
                             if os.path.getsize(dest_path) == os.path.getsize(src_path):
                                if dest_path not in all_copied_to_report_folder:
                                    all_copied_to_report_folder.append(dest_path)
                                if is_local_snapshot:
                                    processed_indices.append(idx)
                                continue
                        shutil.copy2(src_path, dest_path)
                        all_copied_to_report_folder.append(dest_path)
                        copied_count_manual += 1
                        if is_local_snapshot:
                            processed_indices.append(idx)
                    except Exception as e:
                        self.update_status(f"Error copying attachment '{os.path.basename(src_path)}': {e}", True)
                else:
                    if is_local_snapshot:
                        processed_indices.append(idx)

            new_manual_attachments = []
            for i, attachment_path in enumerate(attachments_to_process):
                if i in processed_indices:
                    is_snapshot_for_deletion = CONFIG_DIR in attachment_path and "temp_watched_files" in attachment_path
                    if is_snapshot_for_deletion and os.path.exists(attachment_path):
                        try:
                            os.remove(attachment_path)
                        except Exception as e_del:
                            new_manual_attachments.append(attachment_path)
                else:
                    new_manual_attachments.append(attachment_path)
            self.manual_attachments = new_manual_attachments
            if copied_count_manual > 0:
                self.update_status(f"Processed {copied_count_manual} staged attachment(s).")
            self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")
        return all_copied_to_report_folder

    # --- MAIN SUBMIT FORM ---
    def submit_form(self):
        # ... (Validation logic is unchanged) ...
        patient_name = self.entry_patient_name.get().strip()
        patient_id_input = self.entry_patient_id.get().strip().upper()
        accession_number = self.entry_accession.get().strip().upper()
        dob_input_str = self.entry_dob.get().strip()
        sex = self.entry_sex.get().strip().upper()
        study_description = self.entry_study_description.get().strip()
        referred_from_original_case = self.entry_referred_from.get().strip()
        requesting_physician = self.entry_requesting_physician.get().strip()
        requested_procedure_id = accession_number
        scheduled_station_ae = self.config.get("Preferences", "default_scheduled_station_ae", fallback="ANY_MODALITY")
        referred_from_for_check = referred_from_original_case.upper()

        dob_yyyymmdd = ""; parsed_successfully = False
        if dob_input_str:
            cleaned_dob_input = "".join(filter(str.isdigit, dob_input_str))
            if len(cleaned_dob_input) == 8:
                for fmt_in in ["%d%m%Y", "%Y%m%d", "%m%d%Y"]:
                    try:
                        dt_obj = datetime.strptime(cleaned_dob_input, fmt_in)
                        dob_yyyymmdd = dt_obj.strftime("%Y%m%d")
                        parsed_successfully = True; break
                    except ValueError: continue
            if not parsed_successfully:
                for sep_fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y"]:
                    try:
                        dt_obj = datetime.strptime(dob_input_str, sep_fmt)
                        dob_yyyymmdd = dt_obj.strftime("%Y%m%d")
                        parsed_successfully = True; break
                    except ValueError: continue

        required_fields_map = {
            "patient_name": patient_name, "patient_id": patient_id_input, "accession": accession_number,
            "dob": dob_input_str, "sex": sex, "study_description": study_description,
            "referred_from": referred_from_original_case, "requesting_physician": requesting_physician
        }
        for key, val in required_fields_map.items():
            if not val:
                field_display_name = self.get_ui_label(key, key.replace("_", " ").title())
                messagebox.showerror("Validation Error", f"{field_display_name} is required!", parent=self.root)
                return

        if not dob_yyyymmdd:
            messagebox.showerror("Validation Error", f"Date of Birth '{dob_input_str}' is invalid.", parent=self.root)
            return

        modality = get_modality_from_accession(accession_number)
        if not modality:
            messagebox.showerror("Validation Error", "Invalid Accession Number format.", parent=self.root)
            return

        processed_pid = patient_id_input
        if ' ' not in patient_id_input and len(patient_id_input) >= 8 and patient_id_input[:6].isdigit() and patient_id_input[6:].isalnum():
            processed_pid = f"{patient_id_input[:6]} {patient_id_input[6:]}"
        elif len(processed_pid) < 7:
            messagebox.showerror("Validation Error", "Patient ID is too short or has an invalid format.", parent=self.root)
            return

        if sex not in ['M', 'F', 'O']:
            messagebox.showerror("Validation Error", "Sex must be 'M', 'F', or 'O'!", parent=self.root)
            return

        is_dup, hrs, mins, prev_dt_str, prev_desc = self.check_for_duplicates(patient_name, processed_pid, accession_number)
        if is_dup:
            dup_msg = f"A recent record for Patient '{patient_name}' (or ID '{processed_pid}') was found.\n"
            dup_msg += f"Previous Study: {prev_desc}\nRegistered: ~{hrs}h {mins}m ago ({prev_dt_str}).\n\n"
            dup_msg += "Do you want to register this new study anyway?"
            if not messagebox.askyesno("Duplicate Warning", dup_msg, parent=self.root):
                self.clear_form_fields()
                return

        study_date_now = datetime.now().strftime("%Y%m%d")
        study_time_now = datetime.now().strftime("%H%M%S")

        patient_data = {
            "Patient Name": patient_name, "Patient ID": processed_pid,
            "Accession Number": accession_number, "Date of Birth": dob_yyyymmdd,
            "Sex": sex, "Study Date": study_date_now, "Study Time": study_time_now,
            "Study Description": study_description, "Referred From": referred_from_original_case,
            "Modality": modality, "Requesting Physician": requesting_physician,
            "Requested Procedure ID": requested_procedure_id, "Scheduled Station AE Title": scheduled_station_ae
        }

        try:
            add_patient_record_db(patient_data)
            self.update_status("Patient data saved. Available to MWL Server.")
            self.load_combobox_values_from_db()
        except Exception as e:
            messagebox.showerror("Database Error", f"An unexpected error occurred: {e}", parent=self.root)
            return

        modality_base_output_folder = self.get_modality_specific_path("Paths.Output.DOCX.Modalities", modality)
        if not modality_base_output_folder:
            messagebox.showerror("Configuration Error", f"Base output path for modality '{modality}' is not configured!", parent=self.root)

        safe_pname_folder = "".join(c if c.isalnum() else "_" for c in patient_name)
        safe_pid_folder = processed_pid.replace(' ', '_')
        patient_subfolder_name = f"{safe_pname_folder}_{safe_pid_folder}"
        final_patient_report_folder = os.path.join(modality_base_output_folder, patient_subfolder_name)

        generate_report_and_folder = True
        if modality == "DX":
            referred_ok = any(referred_from_for_check.startswith(p) for p in ["H/", "HE", "PP", "PR"])
            if not (referred_ok or len(self.manual_attachments) > 0):
                generate_report_and_folder = False
                messagebox.showinfo("Partial Registration (DX)",
                                    f"Patient '{patient_name}' registered.\n\n"
                                    f"DOCX Report for {modality} was NOT generated (criteria not met).",
                                    parent=self.root)

        generated_docx_path = None
        if generate_report_and_folder and modality_base_output_folder:
            ensure_dir_exists(final_patient_report_folder)
            generated_docx_path = self.generate_docx_report(patient_data, modality, final_patient_report_folder)
            if generated_docx_path:
                self.process_attachments(processed_pid, modality, final_patient_report_folder)
                # ** START DICOM WORKFLOW **
                self._start_pacs_query_job(patient_data)
            else:
                self.update_status("Patient registered. DOCX report generation failed.", 3000)

        self.config.set("Preferences", "last_referred_from", referred_from_original_case)
        save_config(self.config)

        if generate_report_and_folder and generated_docx_path:
            messagebox.showinfo("Success",
                                f"PATIENT REGISTERED!\n\nPatient: {patient_name} ({processed_pid})\n"
                                f"Report & attachments processed.",
                                parent=self.root)
        self.clear_form_fields()
        self.entry_patient_id.focus()

    # ... (confirm_clear_form, clear_form_fields, on_patient_id_change, populate_fields are unchanged) ...
    def confirm_clear_form(self):
        if messagebox.askyesno("Confirm Clear", "Are you sure you want to clear all fields and staged attachments?", parent=self.root):
            self.clear_form_fields()
            self.update_status("Form cleared.")

    def clear_form_fields(self):
     for attr_name in ['entry_patient_id', 'entry_patient_name', 'entry_dob', 'entry_sex']:
        if hasattr(self, attr_name):
            getattr(self, attr_name).delete(0, tk.END)
     if hasattr(self, 'entry_accession'):
        self.entry_accession.delete(0, tk.END)
        self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))
     for combo_attr in ['entry_study_description', 'entry_referred_from', 'entry_requesting_physician']:
        if hasattr(self, combo_attr) and isinstance(getattr(self, combo_attr), ttk.Combobox):
            getattr(self, combo_attr).set('')
     self.manual_attachments.clear()
     self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} 0")
     if hasattr(self, 'entry_patient_id'):
        self.entry_patient_id.focus()

    def on_patient_id_change(self, event=None):
        if not hasattr(self, 'entry_patient_id'): return
        pid_in = self.entry_patient_id.get().strip().upper()
        if not pid_in: return

        current_cursor_pos = self.entry_patient_id.index(tk.INSERT)
        formatted_pid = pid_in

        if ' ' not in pid_in and len(pid_in) >= 8 and pid_in[:6].isdigit() and pid_in[6:].isalnum():
            formatted_pid = f"{pid_in[:6]} {pid_in[6:]}"
            self.entry_patient_id.delete(0, tk.END)
            self.entry_patient_id.insert(0, formatted_pid)
            try:
                self.entry_patient_id.icursor(current_cursor_pos + 1 if current_cursor_pos >= 6 else current_cursor_pos)
            except tk.TclError: pass

        try: self.entry_patient_id.unbind("<FocusOut>")
        except tk.TclError: pass

        data = get_patient_by_id_db(formatted_pid)
        if data:
            self.populate_fields(data)
            self.update_status(f"Data loaded from DB for Patient ID: {formatted_pid}")
        else:
            for attr in ['entry_patient_name', 'entry_dob', 'entry_sex', 'entry_requesting_physician']:
                if hasattr(self, attr):
                    widget = getattr(self,attr)
                    if isinstance(widget, ttk.Combobox): widget.set("")
                    else: widget.delete(0, tk.END)
            self.update_status(f"No data found in DB for Patient ID: {formatted_pid}")

        try:
            self.root.after(100, lambda: self.entry_patient_id.bind("<FocusOut>", self.on_patient_id_change) if hasattr(self, 'entry_patient_id') and self.entry_patient_id.winfo_exists() else None)
        except tk.TclError: pass

    def populate_fields(self, patient_data):
        if hasattr(self, 'entry_patient_name'):
            self.entry_patient_name.delete(0, tk.END)
            self.entry_patient_name.insert(0, patient_data.get('patient_name', ''))
        if hasattr(self, 'entry_dob'):
            self.entry_dob.delete(0, tk.END)
            dob_yyyymmdd = patient_data.get('dob_yyyymmdd', '')
            if dob_yyyymmdd:
                try:
                    dob_display = datetime.strptime(dob_yyyymmdd, "%Y%m%d").strftime("%d/%m/%Y")
                    self.entry_dob.insert(0, dob_display)
                except ValueError:
                    self.entry_dob.insert(0, dob_yyyymmdd)
        if hasattr(self, 'entry_sex'):
            self.entry_sex.delete(0, tk.END)
            self.entry_sex.insert(0, patient_data.get('sex', ''))
        if patient_data.get('referred_from') and hasattr(self, 'entry_referred_from'):
            self.entry_referred_from.set(patient_data.get('referred_from', ''))
        if patient_data.get('requesting_physician') and hasattr(self, 'entry_requesting_physician'):
            self.entry_requesting_physician.set(patient_data.get('requesting_physician', ''))
        if hasattr(self, 'entry_study_description'):
            self.entry_study_description.set('')
        if hasattr(self, 'entry_accession'):
            self.entry_accession.delete(0, tk.END)
            self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))
            
    # ... (open_settings_window and its many sub-methods are unchanged) ...
    def open_settings_window(self):
        # This extensive method remains the same as the previous version.
        # It handles the creation and logic for the entire settings window.
        logging.info("Opening settings window.")
        settings_win = tk.Toplevel(self.root)
        settings_win.title(self.get_ui_label("settings_window_title", "Settings"))
        settings_win.geometry("950x850")
        settings_win.transient(self.root)
        settings_win.grab_set()
        settings_win.configure(bg=self.current_palette.get("bg", "#F0F0F0"))

        tab_control = ttk.Notebook(settings_win)
        # Create all tabs
        paths_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        modality_paths_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        dicom_network_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        app_mwl_server_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        ext_reports_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        appearance_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        prefs_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        smtp_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        email_recipients_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        email_templates_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        ui_labels_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        
        # Add tabs to control
        tab_control.add(paths_tab, text='General Paths')
        tab_control.add(modality_paths_tab, text='Modality Paths')
        tab_control.add(dicom_network_tab, text='DICOM Network')
        tab_control.add(app_mwl_server_tab, text=self.get_ui_label("app_mwl_server_tab_title", "This App as MWL Server"))
        tab_control.add(ext_reports_tab, text=self.get_ui_label("ext_reports_tab_title", "External Reports"))
        tab_control.add(appearance_tab, text=self.get_ui_label("appearance_tab_title", "Appearance"))
        tab_control.add(prefs_tab, text='Preferences')
        tab_control.add(smtp_tab, text=self.get_ui_label("smtp_settings_tab_title", "SMTP Email"))
        tab_control.add(email_recipients_tab, text=self.get_ui_label("email_recipients_tab_title", "Email Recipients"))
        tab_control.add(email_templates_tab, text=self.get_ui_label("email_templates_tab_title", "Email Templates"))
        tab_control.add(ui_labels_tab, text='UI Labels')
        
        tab_control.pack(expand=1, fill="both", padx=10, pady=10)

        self.settings_entries = {}
        self.ui_label_settings_entries = {}
        self.ext_report_watch_folders_listbox = None
        self.favorite_recipients_listbox = None
        self.email_templates_listbox = None
        self.current_template_widgets = {}
        self.app_mwl_config_widgets = {}

        # Populate General Paths Tab
        ttk.Label(paths_tab, text="General File Paths:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        gp_frame = ttk.Frame(paths_tab, style='Custom.TFrame')
        gp_frame.pack(expand=True, fill=tk.BOTH)
        gp_frame.columnconfigure(1, weight=1)
        general_paths_map = [
            ("Database File:", "Paths", "db_file", False),
            ("DOCX Template File:", "Paths", "docx_template", False),
            ("General DOCX Output Folder:", "Paths", "general_docx_output_folder", True),
            ("General Modality Watch Folder (optional):", "Paths", "general_watch_folder", True)
        ]
        for r, (lbl_text, section, key, is_folder) in enumerate(general_paths_map):
            ttk.Label(gp_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(gp_frame, width=70)
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            if key != "db_file":
                ttk.Button(gp_frame, text="Browse...",
                           command=lambda e=entry, f=is_folder: self.browse_path(e, f, parent=settings_win)
                          ).grid(row=r, column=2, padx=5, pady=7)
            else:
                entry.config(state="readonly")
            self.settings_entries[(section, key)] = entry

        # Populate Modality Paths Tab
        ttk.Label(modality_paths_tab, text="Modality-Specific Paths:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        mp_canvas = tk.Canvas(modality_paths_tab, bg=self.current_palette.get("frame_bg"), highlightthickness=0)
        mp_scrollbar = ttk.Scrollbar(modality_paths_tab, orient="vertical", command=mp_canvas.yview)
        mp_scrollable_frame = ttk.Frame(mp_canvas, style='Custom.TFrame')
        mp_scrollable_frame.bind("<Configure>", lambda e: mp_canvas.configure(scrollregion=mp_canvas.bbox("all")))
        mp_canvas_window = mp_canvas.create_window((0, 0), window=mp_scrollable_frame, anchor="nw")
        mp_canvas.configure(yscrollcommand=mp_scrollbar.set)
        mp_canvas.pack(side="left", fill="both", expand=True)
        mp_scrollbar.pack(side="right", fill="y")
        mp_scrollable_frame.columnconfigure(1, weight=1)
        mp_scrollable_frame.columnconfigure(4, weight=1)
        row_idx = 0
        for mod_code in MODALITIES:
            ttk.Label(mp_scrollable_frame, text=f"{mod_code} DOCX Output:", font=('Helvetica', 10, 'bold')).grid(row=row_idx, column=0, sticky=tk.W, padx=5, pady=3)
            entry_docx = ttk.Entry(mp_scrollable_frame, width=35)
            entry_docx.insert(0, self.config.get("Paths.Output.DOCX.Modalities", mod_code, fallback=""))
            entry_docx.grid(row=row_idx, column=1, sticky=tk.EW, padx=5, pady=3)
            ttk.Button(mp_scrollable_frame, text="...", width=3, command=lambda e=entry_docx: self.browse_path(e, True, parent=settings_win)).grid(row=row_idx, column=2, padx=(0,10), pady=3)
            self.settings_entries[("Paths.Output.DOCX.Modalities", mod_code)] = entry_docx

            ttk.Label(mp_scrollable_frame, text=f"{mod_code} Watch Folder:", font=('Helvetica', 10, 'bold')).grid(row=row_idx, column=3, sticky=tk.W, padx=(10,5), pady=3)
            entry_watch = ttk.Entry(mp_scrollable_frame, width=35)
            entry_watch.insert(0, self.config.get("Paths.WatchFolders.Modalities", mod_code, fallback=""))
            entry_watch.grid(row=row_idx, column=4, sticky=tk.EW, padx=5, pady=3)
            ttk.Button(mp_scrollable_frame, text="...", width=3, command=lambda e=entry_watch: self.browse_path(e, True, parent=settings_win)).grid(row=row_idx, column=5, padx=(0,5), pady=3)
            self.settings_entries[("Paths.WatchFolders.Modalities", mod_code)] = entry_watch
            row_idx += 1
        mp_scrollable_frame.bind("<Configure>", lambda e: mp_canvas.itemconfig(mp_canvas_window, width=e.width))
        
        self._setup_dicom_network_tab(dicom_network_tab, settings_win)
        self._setup_app_mwl_server_tab(app_mwl_server_tab)
        
        # Populate External Reports Tab
        ttk.Label(ext_reports_tab, text=self.get_ui_label("ext_reports_watch_folder_label"), font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        er_top_frame = ttk.Frame(ext_reports_tab, style='Custom.TFrame')
        er_top_frame.pack(expand=False, fill=tk.X, pady=(0,10))
        self.ext_report_watch_folders_listbox = tk.Listbox(er_top_frame, selectmode=tk.SINGLE, height=6, exportselection=False,
                                                            bg=self.current_palette.get("entry_bg"),
                                                            fg=self.current_palette.get("entry_fg"),
                                                            selectbackground=self.current_palette.get("header_fg"),
                                                            selectforeground=self.current_palette.get("button_fg"))
        self.ext_report_watch_folders_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=(0,5))
        self.ext_report_watch_folders_listbox.bind("<<ListboxSelect>>", self._on_ext_report_folder_select)
        er_buttons_frame = ttk.Frame(er_top_frame, style='Custom.TFrame')
        er_buttons_frame.pack(side=tk.RIGHT, fill=tk.Y)
        ttk.Button(er_buttons_frame, text="Add Folder", command=lambda: self._add_ext_report_folder(settings_win)).pack(pady=(0,5), fill=tk.X)
        ttk.Button(er_buttons_frame, text="Remove Selected", command=self._remove_ext_report_folder).pack(fill=tk.X)
        autosend_frame = ttk.LabelFrame(ext_reports_tab, text=self.get_ui_label("ext_report_autosend_group_label", "Automatic Emailing for Selected Folder:"),
                                       style='Custom.TFrame', padding=10)
        autosend_frame.pack(expand=False, fill=tk.X, pady=(10,0))
        self.ext_report_autosend_widgets['enable_var'] = tk.BooleanVar()
        enable_check = ttk.Checkbutton(autosend_frame, text=self.get_ui_label("ext_report_autosend_enable_label", "Enable Auto-Emailing for this Folder"),
                                      variable=self.ext_report_autosend_widgets['enable_var'])
        enable_check.pack(anchor=tk.W, pady=(0,5))
        self.ext_report_autosend_widgets['enable'] = enable_check
        recipients_frame = ttk.Frame(autosend_frame, style='Custom.TFrame')
        recipients_frame.pack(fill=tk.X, pady=(0,5))
        ttk.Label(recipients_frame, text=self.get_ui_label("ext_report_autosend_recipients_label", "Auto-Send Recipient(s):")).pack(side=tk.LEFT)
        self.ext_report_autosend_widgets['recipients'] = ttk.Entry(recipients_frame, width=50)
        self.ext_report_autosend_widgets['recipients'].pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(5,0))
        ttk.Button(recipients_frame, text=self.get_ui_label("ext_report_autosend_add_favorite_button", "Add Favorite"),
                  command=lambda: self._add_favorite_to_autosend_recipients(settings_win)).pack(side=tk.RIGHT, padx=(5,0))
        template_frame = ttk.Frame(autosend_frame, style='Custom.TFrame')
        template_frame.pack(fill=tk.X, pady=(0,5))
        ttk.Label(template_frame, text=self.get_ui_label("ext_report_autosend_template_label", "Auto-Send Email Template:")).pack(side=tk.LEFT)
        self.ext_report_autosend_widgets['template'] = ttk.Combobox(template_frame, width=30, state="readonly")
        self.ext_report_autosend_widgets['template'].pack(side=tk.LEFT, padx=(5,0))
        self.ext_report_autosend_widgets['template'].bind("<<ComboboxSelected>>", self._toggle_autosend_custom_fields_state)
        self.ext_report_autosend_widgets['custom_subject_label'] = ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_custom_subject_label", "Custom Subject:"))
        self.ext_report_autosend_widgets['custom_subject'] = ttk.Entry(autosend_frame, width=80)
        self.ext_report_autosend_widgets['custom_body_label'] = ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_custom_body_label", "Custom Body:"))
        self.ext_report_autosend_widgets['custom_body'] = tk.Text(autosend_frame, height=4, width=80,
                                                                  bg=self.current_palette.get("entry_bg"),
                                                                  fg=self.current_palette.get("entry_fg"),
                                                                  insertbackground=self.current_palette.get("entry_fg"))
        ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_placeholders_label", "Placeholders: {Filename}, {FolderPath}, etc."),
                 foreground="gray").pack(anchor=tk.W, pady=(5,0))
        for folder in self.get_all_external_report_watch_folders():
            self.ext_report_watch_folders_listbox.insert(tk.END, folder)
            
        # ... (The rest of the settings tabs are populated here, unchanged) ...
        # Appearance tab
        ttk.Label(appearance_tab, text="Visual Appearance Settings:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        app_frame = ttk.Frame(appearance_tab, style='Custom.TFrame')
        app_frame.pack(expand=True, fill=tk.BOTH)
        ui_style_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        ui_style_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(ui_style_frame, text=self.get_ui_label("ui_style_engine_label", "UI Style Engine:")).pack(side=tk.LEFT)
        ui_style_combo = ttk.Combobox(ui_style_frame, width=25, state="readonly",
                                     values=["System Default", "Clam (Modern)", "Alt (Modern-ish)", "Default (Classic)", "Classic (Older)"])
        ui_style_combo.set(self.config.get("Preferences", "ui_style", fallback="System Default"))
        ui_style_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "ui_style")] = ui_style_combo
        color_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        color_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(color_frame, text=self.get_ui_label("color_palette_label", "Color Palette:")).pack(side=tk.LEFT)
        color_combo = ttk.Combobox(color_frame, width=25, state="readonly",
                                  values=["Default", "Light Blue", "Dark", "High Contrast", "Mint Green", "Lavender"])
        color_combo.set(self.config.get("Preferences", "color_theme", fallback="Default"))
        color_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "color_theme")] = color_combo
        ui_size_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        ui_size_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(ui_size_frame, text="UI Size:").pack(side=tk.LEFT)
        ui_size_combo = ttk.Combobox(ui_size_frame, width=25, state="readonly",
                                    values=["Very Compact", "Compact", "Default", "Large", "Extra Large"])
        ui_size_combo.set(self.config.get("Preferences", "ui_size", fallback="Default"))
        ui_size_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "ui_size")] = ui_size_combo
        tooltip_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        tooltip_frame.pack(fill=tk.X, pady=(0,10))
        tooltip_var = tk.BooleanVar(value=self.config.getboolean("Preferences", "enable_tooltips", fallback=True))
        ttk.Checkbutton(tooltip_frame, text="Enable Tooltips", variable=tooltip_var).pack(side=tk.LEFT)
        self.settings_entries[("Preferences", "enable_tooltips")] = tooltip_var

        # Preferences tab
        ttk.Label(prefs_tab, text="Application Preferences:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        pref_frame = ttk.Frame(prefs_tab, style='Custom.TFrame')
        pref_frame.pack(expand=True, fill=tk.BOTH)
        pref_frame.columnconfigure(1, weight=1)
        prefs_map = [
            ("Default Accession Prefix:", "Preferences", "default_accession_prefix"),
            ("Default Scheduled Station AE:", "Preferences", "default_scheduled_station_ae")
        ]
        for r, (lbl_text, section, key) in enumerate(prefs_map):
            ttk.Label(pref_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(pref_frame, width=40)
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            self.settings_entries[(section, key)] = entry

        # SMTP tab
        ttk.Label(smtp_tab, text="SMTP Email Configuration:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        smtp_frame = ttk.Frame(smtp_tab, style='Custom.TFrame')
        smtp_frame.pack(expand=True, fill=tk.BOTH)
        smtp_frame.columnconfigure(1, weight=1)
        smtp_map = [
            (self.get_ui_label("smtp_server", "SMTP Server:"), "SMTP", "server"),
            (self.get_ui_label("smtp_port", "SMTP Port:"), "SMTP", "port"),
            (self.get_ui_label("smtp_user", "SMTP Username:"), "SMTP", "user"),
            (self.get_ui_label("smtp_password", "SMTP Password:"), "SMTP", "password"),
            (self.get_ui_label("smtp_sender_email", "Sender Email:"), "SMTP", "sender_email")
        ]
        for r, (lbl_text, section, key) in enumerate(smtp_map):
            ttk.Label(smtp_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(smtp_frame, width=50, show="*" if "password" in key else "")
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            self.settings_entries[(section, key)] = entry
        tls_var = tk.BooleanVar(value=self.config.getboolean("SMTP", "use_tls", fallback=True))
        ttk.Checkbutton(smtp_frame, text=self.get_ui_label("smtp_use_tls", "Use TLS/STARTTLS"), variable=tls_var).grid(row=len(smtp_map), column=0, columnspan=2, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("SMTP", "use_tls")] = tls_var
        ttk.Button(smtp_frame, text=self.get_ui_label("smtp_test_button", "Test SMTP Settings"),
                  command=lambda: self.test_smtp_settings(settings_win)).grid(row=len(smtp_map)+1, column=0, columnspan=2, pady=10)

        # Email Recipients tab
        ttk.Label(email_recipients_tab, text=self.get_ui_label("favorite_recipients_label", "Favorite Email Recipients:"),
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        rec_frame = ttk.Frame(email_recipients_tab, style='Custom.TFrame')
        rec_frame.pack(expand=True, fill=tk.BOTH)
        self.favorite_recipients_listbox = tk.Listbox(rec_frame, selectmode=tk.SINGLE, height=10,
                                                     bg=self.current_palette.get("entry_bg"),
                                                     fg=self.current_palette.get("entry_fg"),
                                                     selectbackground=self.current_palette.get("header_fg"),
                                                     selectforeground=self.current_palette.get("button_fg"))
        self.favorite_recipients_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=(0,5))
        rec_buttons_frame = ttk.Frame(rec_frame, style='Custom.TFrame')
        rec_buttons_frame.pack(side=tk.RIGHT, fill=tk.Y)
        ttk.Button(rec_buttons_frame, text=self.get_ui_label("add_favorite_button", "Add Favorite"),
                  command=lambda: self._add_favorite_recipient(settings_win)).pack(pady=(0,5), fill=tk.X)
        ttk.Button(rec_buttons_frame, text=self.get_ui_label("remove_favorite_button", "Remove Selected"),
                  command=self._remove_favorite_recipient).pack(fill=tk.X)
        max_recent_frame = ttk.Frame(email_recipients_tab, style='Custom.TFrame')
        max_recent_frame.pack(fill=tk.X, pady=(10,0))
        ttk.Label(max_recent_frame, text=self.get_ui_label("max_recent_recipients_label", "Max Recent Recipients:")).pack(side=tk.LEFT)
        max_recent_entry = ttk.Entry(max_recent_frame, width=10)
        max_recent_entry.insert(0, self.config.get("EmailRecipients", "max_recent", fallback="10"))
        max_recent_entry.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("EmailRecipients", "max_recent")] = max_recent_entry
        favorites_str = self.config.get("EmailRecipients", "favorite_list", fallback="")
        for email in [e.strip() for e in favorites_str.split(';') if e.strip()]:
            self.favorite_recipients_listbox.insert(tk.END, email)

        # Email Templates tab
        ttk.Label(email_templates_tab, text=self.get_ui_label("manage_email_templates_label", "Manage Email Templates:"),
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        templates_main_frame = ttk.Frame(email_templates_tab, style='Custom.TFrame')
        templates_main_frame.pack(expand=True, fill=tk.BOTH)
        templates_list_frame = ttk.Frame(templates_main_frame, style='Custom.TFrame')
        templates_list_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0,10))
        self.email_templates_listbox = tk.Listbox(templates_list_frame, selectmode=tk.SINGLE, width=25, height=8,
                                                 bg=self.current_palette.get("entry_bg"),
                                                 fg=self.current_palette.get("entry_fg"),
                                                 selectbackground=self.current_palette.get("header_fg"),
                                                 selectforeground=self.current_palette.get("button_fg"))
        self.email_templates_listbox.pack(expand=True, fill=tk.BOTH, pady=(0,5))
        self.email_templates_listbox.bind("<<ListboxSelect>>", lambda e: self._load_template_for_editing())
        temp_buttons_frame = ttk.Frame(templates_list_frame, style='Custom.TFrame')
        temp_buttons_frame.pack(fill=tk.X)
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("add_template_button", "Add New"),
                  command=lambda: self._edit_email_template(settings_win, is_new=True)).pack(fill=tk.X, pady=(0,2))
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("edit_template_button", "Edit Selected"),
                  command=lambda: self._edit_email_template(settings_win, is_new=False)).pack(fill=tk.X, pady=(0,2))
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("delete_template_button", "Delete Selected"),
                  command=self._delete_email_template).pack(fill=tk.X)
        templates_edit_frame = ttk.Frame(templates_main_frame, style='Custom.TFrame')
        templates_edit_frame.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_name_label", "Template Name:")).pack(anchor=tk.W)
        self.current_template_widgets['name'] = ttk.Entry(templates_edit_frame, width=40, state="readonly")
        self.current_template_widgets['name'].pack(fill=tk.X, pady=(0,5))
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_subject_label", "Subject Template:")).pack(anchor=tk.W)
        self.current_template_widgets['subject'] = ttk.Entry(templates_edit_frame, width=80, state="readonly")
        self.current_template_widgets['subject'].pack(fill=tk.X, pady=(0,5))
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_body_label", "Body Template:")).pack(anchor=tk.W)
        self.current_template_widgets['body'] = tk.Text(templates_edit_frame, height=12, width=80, state="disabled",
                                                       bg=self.current_palette.get("entry_bg"),
                                                       fg=self.current_palette.get("entry_fg"))
        self.current_template_widgets['body'].pack(expand=True, fill=tk.BOTH, pady=(0,5))
        temp_edit_buttons_frame = ttk.Frame(templates_edit_frame, style='Custom.TFrame')
        temp_edit_buttons_frame.pack(fill=tk.X)
        ttk.Button(temp_edit_buttons_frame, text="Save Changes", command=self._save_current_template_changes).pack(side=tk.LEFT, padx=(0,5))
        ttk.Button(temp_edit_buttons_frame, text="Clear", command=self._clear_template_editor_fields).pack(side=tk.LEFT)
        ttk.Label(templates_edit_frame, text=self.get_ui_label("available_placeholders_label", "Placeholders: {Patient Name}, etc."),
                 foreground="gray").pack(anchor=tk.W, pady=(5,0))
        self._populate_email_templates_listbox()

        # UI Labels tab
        ttk.Label(ui_labels_tab, text="UI Label Customization:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        ui_canvas = tk.Canvas(ui_labels_tab, bg=self.current_palette.get("frame_bg"), highlightthickness=0)
        ui_scrollbar = ttk.Scrollbar(ui_labels_tab, orient="vertical", command=ui_canvas.yview)
        ui_scrollable_frame = ttk.Frame(ui_canvas, style='Custom.TFrame')
        ui_scrollable_frame.bind("<Configure>", lambda e: ui_canvas.configure(scrollregion=ui_canvas.bbox("all")))
        ui_canvas_window = ui_canvas.create_window((0, 0), window=ui_scrollable_frame, anchor="nw")
        ui_canvas.configure(yscrollcommand=ui_scrollbar.set)
        ui_canvas.pack(side="left", fill="both", expand=True)
        ui_scrollbar.pack(side="right", fill="y")
        ui_scrollable_frame.columnconfigure(1, weight=1)
        ui_labels_items = list(self.config.items("UI.Labels"))
        for r, (key, current_value) in enumerate(ui_labels_items):
            ttk.Label(ui_scrollable_frame, text=f"{key}:").grid(row=r, column=0, sticky=tk.W, padx=5, pady=3)
            entry = ttk.Entry(ui_scrollable_frame, width=60)
            entry.insert(0, current_value)
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=3)
            self.ui_label_settings_entries[key] = entry
        ui_scrollable_frame.bind("<Configure>", lambda e: ui_canvas.itemconfig(ui_canvas_window, width=e.width))
        
        # Save Button
        button_frame = ttk.Frame(settings_win, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)
        ttk.Button(button_frame, text="Save Settings", command=lambda: self.save_settings_and_apply(settings_win)).pack(side=tk.RIGHT, padx=(5,0))
        ttk.Button(button_frame, text="Cancel", command=settings_win.destroy).pack(side=tk.RIGHT)
        
        # Initial population of auto-send template combobox
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        template_names = [t.strip() for t in template_names if t.strip()]
        if template_names:
            template_names.append("Custom")
            self.ext_report_autosend_widgets['template']['values'] = template_names
        self._set_autosend_widgets_state("disabled")

    def save_settings_and_apply(self, settings_win):
        # This comprehensive method handles saving all settings from the UI to the config file
        # and then applying them to the running application. It's unchanged from the previous version.
        logging.info("Saving settings changes.")
        for (section, key), widget in self.settings_entries.items():
            if isinstance(widget, tk.BooleanVar):
                self.config.set(section, key, str(widget.get()))
            elif hasattr(widget, 'get'):
                self.config.set(section, key, widget.get())
        for key, widget in self.ui_label_settings_entries.items():
            self.config.set("UI.Labels", key, widget.get())
        folders = [self.ext_report_watch_folders_listbox.get(i) for i in range(self.ext_report_watch_folders_listbox.size())]
        self.config.set("Paths", "external_report_watch_folders_list", ';'.join(folders))
        for folder in folders:
            folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(folder)}"
            if not self.config.has_section(folder_section):
                self.config.add_section(folder_section)
            if (self._selected_ext_report_folder_for_autosend_config and
                self._selected_ext_report_folder_for_autosend_config == folder):
                self.config.set(folder_section, "enabled", str(self.ext_report_autosend_widgets['enable_var'].get()))
                self.config.set(folder_section, "recipients", self.ext_report_autosend_widgets['recipients'].get())
                self.config.set(folder_section, "template", self.ext_report_autosend_widgets['template'].get())
                self.config.set(folder_section, "custom_subject", self.ext_report_autosend_widgets['custom_subject'].get())
                self.config.set(folder_section, "custom_body", self.ext_report_autosend_widgets['custom_body'].get("1.0", tk.END).strip())
        favorites = [self.favorite_recipients_listbox.get(i) for i in range(self.favorite_recipients_listbox.size())]
        self.config.set("EmailRecipients", "favorite_list", ';'.join(favorites))
        save_config(self.config)
        self.apply_theme_and_styles()
        self.apply_ui_size()
        self.load_combobox_values_from_db()
        self._start_realtime_watchers()
        if PYNETDICOM_AVAILABLE:
            mwl_enabled = self.app_mwl_config_widgets['enabled_var'].get()
            current_mwl_enabled = self.config.getboolean("MWLServerConfig", "enabled", fallback=False)
            if mwl_enabled != current_mwl_enabled:
                if mwl_enabled:
                    self.start_mwl_server_if_configured()
                else:
                    self.stop_mwl_server()
        messagebox.showinfo("Settings Saved", "Settings have been saved and applied.", parent=settings_win)
        settings_win.destroy()
        
    # ... (all other settings helper methods are unchanged) ...
    def _setup_app_mwl_server_tab(self, tab):
        ttk.Label(tab, text=self.get_ui_label("app_mwl_server_tab_title", "MWL Server Configuration:"),
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        if not PYNETDICOM_AVAILABLE:
            ttk.Label(tab,
                     text="⚠️ MWL Server functionality is disabled because pynetdicom/pydicom libraries are not available.",
                     foreground="red", font=('Helvetica', 11, 'bold')).pack(anchor=tk.W)
        mwl_frame = ttk.Frame(tab, style='Custom.TFrame')
        mwl_frame.pack(expand=True, fill=tk.BOTH)
        mwl_frame.columnconfigure(1, weight=1)
        self.app_mwl_config_widgets['enabled_var'] = tk.BooleanVar(value=self.config.getboolean("MWLServerConfig", "enabled", fallback=False))
        enabled_check = ttk.Checkbutton(mwl_frame, text=self.get_ui_label("app_mwl_server_enable_label", "Enable MWL Server"),
                                       variable=self.app_mwl_config_widgets['enabled_var'],
                                       state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        enabled_check.grid(row=0, column=0, columnspan=2, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "enabled")] = self.app_mwl_config_widgets['enabled_var']
        ttk.Label(mwl_frame, text=self.get_ui_label("app_mwl_server_ae_title_label", "AE Title:")).grid(row=1, column=0, sticky=tk.W, padx=5, pady=7)
        ae_entry = ttk.Entry(mwl_frame, width=20, state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        ae_entry.insert(0, self.config.get("MWLServerConfig", "ae_title", fallback="PYREGMWL"))
        ae_entry.grid(row=1, column=1, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "ae_title")] = ae_entry
        ttk.Label(mwl_frame, text=self.get_ui_label("app_mwl_server_port_label", "Port:")).grid(row=2, column=0, sticky=tk.W, padx=5, pady=7)
        port_entry = ttk.Entry(mwl_frame, width=10, state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        port_entry.insert(0, self.config.get("MWLServerConfig", "port", fallback="11112"))
        port_entry.grid(row=2, column=1, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "port")] = port_entry
        
    def _setup_dicom_network_tab(self, tab, parent_window):
        if not PYNETDICOM_AVAILABLE:
            ttk.Label(tab, text="⚠️ DICOM client functionality is disabled (pynetdicom/pydicom not found).",
                      foreground="red", font=('Helvetica', 11, 'bold')).pack(anchor=tk.W)
        pacs_frame = ttk.LabelFrame(tab, text="Primary Query PACS (for finding Study UID)", style='Custom.TFrame', padding=10)
        pacs_frame.pack(fill=tk.X, pady=(5, 15))
        pacs_frame.columnconfigure(1, weight=1)
        pacs_ae_entry = ttk.Entry(pacs_frame)
        pacs_ip_entry = ttk.Entry(pacs_frame)
        pacs_port_entry = ttk.Entry(pacs_frame)
        ttk.Label(pacs_frame, text="AE Title:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        pacs_ae_entry.insert(0, self.config.get("DICOM.QueryPACS", "ae_title", fallback=""))
        pacs_ae_entry.grid(row=0, column=1, sticky=tk.EW, padx=5, pady=5)
        self.settings_entries[("DICOM.QueryPACS", "ae_title")] = pacs_ae_entry
        ttk.Label(pacs_frame, text="IP Address:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        pacs_ip_entry.insert(0, self.config.get("DICOM.QueryPACS", "ip", fallback=""))
        pacs_ip_entry.grid(row=1, column=1, sticky=tk.EW, padx=5, pady=5)
        self.settings_entries[("DICOM.QueryPACS", "ip")] = pacs_ip_entry
        ttk.Label(pacs_frame, text="Port:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        pacs_port_entry.insert(0, self.config.get("DICOM.QueryPACS", "port", fallback=""))
        pacs_port_entry.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        self.settings_entries[("DICOM.QueryPACS", "port")] = pacs_port_entry
        pacs_test_button = ttk.Button(pacs_frame, text="Test/Verify",
            command=lambda: self._test_dicom_echo(
                pacs_ae_entry.get(), pacs_ip_entry.get(), pacs_port_entry.get(), parent_window
            )
        )
        pacs_test_button.grid(row=2, column=2, padx=10, pady=5)
        dest_lf = ttk.LabelFrame(tab, text="Modality DICOM Destinations (for sending converted files)", style='Custom.TFrame', padding=10)
        dest_lf.pack(expand=True, fill=tk.BOTH)
        dest_canvas = tk.Canvas(dest_lf, bg=self.current_palette.get("frame_bg"), highlightthickness=0)
        dest_scrollbar = ttk.Scrollbar(dest_lf, orient="vertical", command=dest_canvas.yview)
        dest_scrollable_frame = ttk.Frame(dest_canvas, style='Custom.TFrame')
        dest_scrollable_frame.bind("<Configure>", lambda e: dest_canvas.configure(scrollregion=dest_canvas.bbox("all")))
        dest_canvas_window = dest_canvas.create_window((0, 0), window=dest_scrollable_frame, anchor="nw")
        dest_canvas.configure(yscrollcommand=dest_scrollbar.set)
        dest_canvas.pack(side="left", fill="both", expand=True)
        dest_scrollbar.pack(side="right", fill="y")
        dest_scrollable_frame.columnconfigure(1, weight=1)
        dest_scrollable_frame.columnconfigure(3, weight=1)
        dest_scrollable_frame.columnconfigure(5, weight=1)
        ttk.Label(dest_scrollable_frame, text="Modality", font=('Helvetica', 10, 'bold')).grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(dest_scrollable_frame, text="Destination AE Title", font=('Helvetica', 10, 'bold')).grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(dest_scrollable_frame, text="IP Address", font=('Helvetica', 10, 'bold')).grid(row=0, column=3, padx=5, pady=5)
        ttk.Label(dest_scrollable_frame, text="Port", font=('Helvetica', 10, 'bold')).grid(row=0, column=5, padx=5, pady=5)
        for r, mod_code in enumerate(MODALITIES, start=1):
            section = f"DICOM.Destinations.{mod_code}"
            ttk.Label(dest_scrollable_frame, text=f"{mod_code}:", font=('Helvetica', 10, 'bold')).grid(row=r, column=0, sticky=tk.W, padx=5, pady=5)
            ae_entry = ttk.Entry(dest_scrollable_frame)
            ae_entry.insert(0, self.config.get(section, "ae_title", fallback=""))
            ae_entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=5)
            self.settings_entries[(section, "ae_title")] = ae_entry
            ip_entry = ttk.Entry(dest_scrollable_frame)
            ip_entry.insert(0, self.config.get(section, "ip", fallback=""))
            ip_entry.grid(row=r, column=3, sticky=tk.EW, padx=5, pady=5)
            self.settings_entries[(section, "ip")] = ip_entry
            port_entry = ttk.Entry(dest_scrollable_frame, width=8)
            port_entry.insert(0, self.config.get(section, "port", fallback=""))
            port_entry.grid(row=r, column=5, sticky=tk.W, padx=5, pady=5)
            self.settings_entries[(section, "port")] = port_entry
            test_button = ttk.Button(dest_scrollable_frame, text="Test", width=5,
                command=lambda ae=ae_entry, ip=ip_entry, p=port_entry: self._test_dicom_echo(
                    ae.get(), ip.get(), p.get(), parent_window
                )
            )
            test_button.grid(row=r, column=6, padx=5, pady=5)
        dest_scrollable_frame.bind("<Configure>", lambda e: dest_canvas.itemconfig(dest_canvas_window, width=e.width))
        if not PYNETDICOM_AVAILABLE:
            pacs_test_button.config(state="disabled")
            for child in dest_scrollable_frame.winfo_children():
                if isinstance(child, ttk.Button):
                    child.config(state="disabled")
                    
    def _set_autosend_widgets_state(self, state):
        for widget_key in ['enable', 'recipients', 'template']:
            if widget_key in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets[widget_key].config(state=state)
        self._toggle_autosend_custom_fields_state()

    def _toggle_autosend_custom_fields_state(self, event=None):
        selected_template = self.ext_report_autosend_widgets.get('template', {}).get() if 'template' in self.ext_report_autosend_widgets else ""
        if selected_template == "Custom":
            if 'custom_subject_label' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_subject_label'].pack(anchor=tk.W, pady=(5,0))
            if 'custom_subject' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_subject'].pack(fill=tk.X, pady=(0,5))
            if 'custom_body_label' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_body_label'].pack(anchor=tk.W, pady=(5,0))
            if 'custom_body' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_body'].pack(fill=tk.BOTH, expand=True, pady=(0,5))
        else:
            for widget_key in ['custom_subject_label', 'custom_subject', 'custom_body_label', 'custom_body']:
                if widget_key in self.ext_report_autosend_widgets:
                    self.ext_report_autosend_widgets[widget_key].pack_forget()
                    
    def _add_favorite_to_autosend_recipients(self, parent_win):
        favorites_str = self.config.get("EmailRecipients", "favorite_list", fallback="")
        favorites = [e.strip() for e in favorites_str.split(';') if e.strip()]
        if not favorites:
            messagebox.showinfo("No Favorites", "No favorite recipients configured.", parent=parent_win)
            return
        def on_add():
            selected = favorite_listbox.curselection()
            if selected:
                email = favorite_listbox.get(selected[0])
                current_recipients = self.ext_report_autosend_widgets['recipients'].get()
                new_recipients = f"{current_recipients}; {email}" if current_recipients else email
                self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
                self.ext_report_autosend_widgets['recipients'].insert(0, new_recipients)
                add_win.destroy()
        add_win = tk.Toplevel(parent_win)
        add_win.title("Select Favorite Recipient")
        add_win.geometry("400x300")
        add_win.transient(parent_win)
        add_win.grab_set()
        ttk.Label(add_win, text="Select a favorite recipient to add:").pack(pady=10)
        favorite_listbox = tk.Listbox(add_win, selectmode=tk.SINGLE)
        favorite_listbox.pack(expand=True, fill=tk.BOTH, padx=10, pady=(0,10))
        for email in favorites:
            favorite_listbox.insert(tk.END, email)
        button_frame = ttk.Frame(add_win)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="Add", command=on_add).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=add_win.destroy).pack(side=tk.LEFT, padx=5)

    def _on_ext_report_folder_select(self, event=None):
        selection = self.ext_report_watch_folders_listbox.curselection()
        if not selection:
            self._set_autosend_widgets_state("disabled")
            self._selected_ext_report_folder_for_autosend_config = None
            return
        selected_folder = self.ext_report_watch_folders_listbox.get(selection[0])
        self._selected_ext_report_folder_for_autosend_config = selected_folder
        self._set_autosend_widgets_state("normal")
        folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(selected_folder)}"
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        template_names = [t.strip() for t in template_names if t.strip()]
        template_names.append("Custom")
        self.ext_report_autosend_widgets['template']['values'] = template_names
        if self.config.has_section(folder_section):
            enabled = self.config.getboolean(folder_section, "enabled", fallback=False)
            recipients = self.config.get(folder_section, "recipients", fallback="")
            template = self.config.get(folder_section, "template", fallback="")
            custom_subject = self.config.get(folder_section, "custom_subject", fallback="")
            custom_body = self.config.get(folder_section, "custom_body", fallback="")
            self.ext_report_autosend_widgets['enable_var'].set(enabled)
            self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
            self.ext_report_autosend_widgets['recipients'].insert(0, recipients)
            if template in template_names:
                self.ext_report_autosend_widgets['template'].set(template)
            else:
                self.ext_report_autosend_widgets['template'].set("")
            self.ext_report_autosend_widgets['custom_subject'].delete(0, tk.END)
            self.ext_report_autosend_widgets['custom_subject'].insert(0, custom_subject)
            self.ext_report_autosend_widgets['custom_body'].delete("1.0", tk.END)
            self.ext_report_autosend_widgets['custom_body'].insert("1.0", custom_body)
        else:
            self.ext_report_autosend_widgets['enable_var'].set(False)
            self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
            self.ext_report_autosend_widgets['template'].set("")
            self.ext_report_autosend_widgets['custom_subject'].delete(0, tk.END)
            self.ext_report_autosend_widgets['custom_body'].delete("1.0", tk.END)
        self._toggle_autosend_custom_fields_state()

    def _add_ext_report_folder(self, parent_win):
        folder = filedialog.askdirectory(title="Select External Report Watch Folder", parent=parent_win)
        if folder:
            folder = os.path.normpath(folder)
            existing_folders = [self.ext_report_watch_folders_listbox.get(i) for i in range(self.ext_report_watch_folders_listbox.size())]
            if folder not in existing_folders:
                self.ext_report_watch_folders_listbox.insert(tk.END, folder)
            else:
                messagebox.showinfo("Duplicate Folder", "This folder is already in the list.", parent=parent_win)

    def _remove_ext_report_folder(self):
        selection = self.ext_report_watch_folders_listbox.curselection()
        if selection:
            folder = self.ext_report_watch_folders_listbox.get(selection[0])
            self.ext_report_watch_folders_listbox.delete(selection[0])
            folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(folder)}"
            if self.config.has_section(folder_section):
                self.config.remove_section(folder_section)
            self._set_autosend_widgets_state("disabled")
            self._selected_ext_report_folder_for_autosend_config = None

    def _add_favorite_recipient(self, parent_win):
        email = simpledialog.askstring("Add Favorite Recipient", "Enter email address:", parent=parent_win)
        if email and email.strip() and "@" in email and "." in email.split("@")[-1]:
            existing = [self.favorite_recipients_listbox.get(i) for i in range(self.favorite_recipients_listbox.size())]
            if email not in existing:
                self.favorite_recipients_listbox.insert(tk.END, email)
            else:
                messagebox.showinfo("Duplicate Email", "This email is already in the favorites list.", parent=parent_win)
        elif email:
            messagebox.showerror("Invalid Email", "Please enter a valid email address.", parent=parent_win)

    def _remove_favorite_recipient(self):
        selection = self.favorite_recipients_listbox.curselection()
        if selection:
            self.favorite_recipients_listbox.delete(selection[0])

    def _populate_email_templates_listbox(self):
        self.email_templates_listbox.delete(0, tk.END)
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        for template in [t.strip() for t in template_names if t.strip()]:
            self.email_templates_listbox.insert(tk.END, template)

    def _clear_template_editor_fields(self):
        for widget in self.current_template_widgets.values():
            if isinstance(widget, ttk.Entry):
                widget.config(state="normal")
                widget.delete(0, tk.END)
                widget.config(state="readonly")
            elif isinstance(widget, tk.Text):
                widget.config(state="normal")
                widget.delete("1.0", tk.END)
                widget.config(state="disabled")

    def _load_template_for_editing(self, template_name_to_load=None):
        if template_name_to_load is None:
            selection = self.email_templates_listbox.curselection()
            if not selection:
                self._clear_template_editor_fields()
                return
            template_name_to_load = self.email_templates_listbox.get(selection[0])
        self._selected_template_for_edit = template_name_to_load
        subject = self.config.get("EmailTemplates", f"{template_name_to_load}_subject", fallback="")
        body = self.config.get("EmailTemplates", f"{template_name_to_load}_body", fallback="")
        self.current_template_widgets['name'].config(state="normal")
        self.current_template_widgets['name'].delete(0, tk.END)
        self.current_template_widgets['name'].insert(0, template_name_to_load)
        self.current_template_widgets['name'].config(state="readonly")
        self.current_template_widgets['subject'].config(state="normal")
        self.current_template_widgets['subject'].delete(0, tk.END)
        self.current_template_widgets['subject'].insert(0, subject)
        self.current_template_widgets['body'].config(state="normal")
        self.current_template_widgets['body'].delete("1.0", tk.END)
        self.current_template_widgets['body'].insert("1.0", body)

    def _edit_email_template(self, parent_win, is_new=False):
        if is_new:
            template_name = simpledialog.askstring("New Template", "Enter template name:", parent=parent_win)
            if not template_name or not template_name.strip(): return
            template_name = template_name.strip()
            existing_templates = [t.strip() for t in self.config.get("EmailTemplates", "template_names", fallback="").split(';') if t.strip()]
            if template_name in existing_templates:
                messagebox.showerror("Duplicate Template", "A template with this name already exists.", parent=parent_win)
                return
            new_template_list = ';'.join(existing_templates + [template_name]) if existing_templates else template_name
            self.config.set("EmailTemplates", "template_names", new_template_list)
            self.config.set("EmailTemplates", f"{template_name}_subject", "Report: {Patient Name}")
            self.config.set("EmailTemplates", f"{template_name}_body", "Please find attached the report for {Patient Name}.")
            self._populate_email_templates_listbox()
            for i in range(self.email_templates_listbox.size()):
                if self.email_templates_listbox.get(i) == template_name:
                    self.email_templates_listbox.selection_set(i)
                    break
            self._load_template_for_editing(template_name)
        if hasattr(self, '_selected_template_for_edit') and self._selected_template_for_edit:
            self.current_template_widgets['subject'].config(state="normal")
            self.current_template_widgets['body'].config(state="normal")

    def _save_current_template_changes(self):
        if not hasattr(self, '_selected_template_for_edit') or not self._selected_template_for_edit:
            messagebox.showwarning("No Template Selected", "Please select a template to edit.")
            return
        template_name = self._selected_template_for_edit
        subject = self.current_template_widgets['subject'].get()
        body = self.current_template_widgets['body'].get("1.0", tk.END).strip()
        self.config.set("EmailTemplates", f"{template_name}_subject", subject)
        self.config.set("EmailTemplates", f"{template_name}_body", body)
        self.current_template_widgets['subject'].config(state="readonly")
        self.current_template_widgets['body'].config(state="disabled")
        messagebox.showinfo("Template Saved", f"Template '{template_name}' has been saved.")

    def _delete_email_template(self):
        selection = self.email_templates_listbox.curselection()
        if not selection: return
        template_name = self.email_templates_listbox.get(selection[0])
        if template_name == "Default":
            messagebox.showerror("Cannot Delete", "The Default template cannot be deleted.")
            return
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete the template '{template_name}'?"):
            existing_templates = [t.strip() for t in self.config.get("EmailTemplates", "template_names", fallback="").split(';') if t.strip() and t != template_name]
            self.config.set("EmailTemplates", "template_names", ';'.join(existing_templates))
            self.config.remove_option("EmailTemplates", f"{template_name}_subject")
            self.config.remove_option("EmailTemplates", f"{template_name}_body")
            self._populate_email_templates_listbox()
            self._clear_template_editor_fields()
            self._selected_template_for_edit = None
            
    def browse_path(self, entry_widget, is_folder, parent=None):
        path = filedialog.askdirectory(title="Select Folder", parent=parent or self.root) if is_folder else filedialog.askopenfilename(title="Select File", parent=parent or self.root)
        if path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, os.path.normpath(path))

    # ... (view_patient_data_window, sort_treeview_column, update_email_button_state, trigger_email_report_picker, etc. are unchanged) ...
    def view_patient_data_window(self):
        logging.info("Opening patient data viewer window.")
        data_win = tk.Toplevel(self.root)
        data_win.title(self.get_ui_label("view_data_window_title", "Patient Data Viewer"))
        data_win.geometry("1200x700")
        data_win.transient(self.root)
        data_win.configure(bg=self.current_palette.get("bg", "#F0F0F0"))
        search_frame = ttk.Frame(data_win, style='Custom.TFrame', padding=10)
        search_frame.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(search_frame, text="Search:").pack(side=tk.LEFT, padx=(0,5))
        search_entry = ttk.Entry(search_frame, width=30)
        search_entry.pack(side=tk.LEFT, padx=(0,10))
        tree_frame = ttk.Frame(data_win)
        tree_frame.pack(expand=True, fill=tk.BOTH, padx=10, pady=(0,10))
        columns = ("id", "patient_name", "patient_id", "accession_number", "dob_yyyymmdd", "sex",
                  "study_date", "study_time", "study_description", "referred_from", "modality",
                  "requesting_physician", "requested_procedure_id", "scheduled_station_ae_title", "created_at")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20)
        column_widths = {"id": 50, "patient_name": 150, "patient_id": 100, "accession_number": 120, "dob_yyyymmdd": 90,
                        "sex": 40, "study_date": 80, "study_time": 70, "study_description": 200, "referred_from": 120,
                        "modality": 70, "requesting_physician": 150, "requested_procedure_id": 120, "scheduled_station_ae_title": 100, "created_at": 130}
        for col in columns:
            tree.heading(col, text=col.replace("_", " ").title(), command=lambda c=col: self.sort_treeview_column(tree, c, False))
            tree.column(col, width=column_widths.get(col, 100), minwidth=50)
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=tree_scroll.set)
        tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        def filter_treeview_data(event=None):
            search_term = search_entry.get().strip()
            tree.delete(*tree.get_children())
            try:
                all_data, _ = get_all_patient_records_db(search_term)
                for record in all_data:
                    values = [record.get(col, "") for col in columns]
                    tree.insert("", tk.END, values=values)
            except Exception as e:
                messagebox.showerror("Database Error", f"Error loading patient data: {e}", parent=data_win)
        search_entry.bind("<KeyRelease>", filter_treeview_data)
        ttk.Button(search_frame, text="Refresh", command=filter_treeview_data).pack(side=tk.LEFT, padx=(10,0))
        filter_treeview_data()

    def sort_treeview_column(self, tv, col, reverse):
        data_list = [(tv.set(k, col), k) for k in tv.get_children('')]
        try:
            data_list.sort(key=lambda x: float(x[0]) if x[0].replace('.', '', 1).isdigit() else x[0].lower(), reverse=reverse)
        except (ValueError, AttributeError):
            data_list.sort(key=lambda x: str(x[0]).lower(), reverse=reverse)
        for index, (val, k) in enumerate(data_list):
            tv.move(k, '', index)
        tv.heading(col, command=lambda: self.sort_treeview_column(tv, col, not reverse))

    def update_email_button_state(self):
        try:
            all_data, _ = get_all_patient_records_db("")
            self.email_button.config(state=tk.NORMAL if all_data else tk.DISABLED)
        except Exception:
            self.email_button.config(state=tk.DISABLED)

    def trigger_email_report_picker(self):
        ReportPickerDialog(self.root, self.config, self.current_palette, self.get_ui_label, self)

    def update_recent_recipients(self, recipient_email):
        current_recents_str = self.config.get("EmailRecipients", "recent_list", fallback="")
        current_recents = [e.strip() for e in current_recents_str.split(';') if e.strip()]
        if recipient_email in current_recents:
            current_recents.remove(recipient_email)
        current_recents.insert(0, recipient_email)
        max_recent = self.config.getint("EmailRecipients", "max_recent", fallback=10)
        self.config.set("EmailRecipients", "recent_list", ';'.join(current_recents[:max_recent]))
        save_config(self.config)

    def send_email_with_report(self, recipient_emails_str, subject, body, attachment_paths, update_recents=True):
        smtp_server = self.config.get("SMTP", "server", fallback="")
        smtp_port = self.config.getint("SMTP", "port", fallback=587)
        smtp_user = self.config.get("SMTP", "user", fallback="")
        smtp_password = self.config.get("SMTP", "password", fallback="")
        sender_email = self.config.get("SMTP", "sender_email", fallback="")
        use_tls = self.config.getboolean("SMTP", "use_tls", fallback=True)
        if not smtp_server or not sender_email:
            messagebox.showerror("SMTP Configuration Error", "SMTP server and sender email must be configured.", parent=self.root)
            return False
        try:
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = recipient_emails_str
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))
            for attachment_path in attachment_paths:
                if os.path.isfile(attachment_path):
                    with open(attachment_path, "rb") as attachment:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', f'attachment; filename= {os.path.basename(attachment_path)}')
                        msg.attach(part)
            server = smtplib.SMTP(smtp_server, smtp_port)
            if use_tls:
                server.starttls()
            if smtp_user and smtp_password:
                server.login(smtp_user, smtp_password)
            server.sendmail(sender_email, [e.strip() for e in recipient_emails_str.split(';')], msg.as_string())
            server.quit()
            if update_recents:
                for recipient in [e.strip() for e in recipient_emails_str.split(';') if e.strip()]:
                    self.update_recent_recipients(recipient)
            self.update_status(f"Email sent successfully to: {recipient_emails_str}")
            return True
        except Exception as e:
            messagebox.showerror("Email Error", f"Error sending email: {e}", parent=self.root)
            return False

    def test_smtp_settings(self, parent_window):
        smtp_server = self.settings_entries[("SMTP", "server")].get()
        smtp_port_str = self.settings_entries[("SMTP", "port")].get()
        smtp_user = self.settings_entries[("SMTP", "user")].get()
        smtp_password = self.settings_entries[("SMTP", "password")].get()
        sender_email = self.settings_entries[("SMTP", "sender_email")].get()
        use_tls = self.settings_entries[("SMTP", "use_tls")].get()
        if not smtp_server or not sender_email:
            messagebox.showerror("Configuration Error", "SMTP server and sender email are required for testing.", parent=parent_window)
            return
        try:
            smtp_port = int(smtp_port_str)
        except ValueError:
            messagebox.showerror("Configuration Error", "SMTP port must be a valid number.", parent=parent_window)
            return
        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            if use_tls:
                server.starttls()
            if smtp_user and smtp_password:
                server.login(smtp_user, smtp_password)
            server.quit()
            messagebox.showinfo("SMTP Test Successful", "SMTP settings are working correctly!", parent=parent_window)
        except Exception as e:
            messagebox.showerror("SMTP Test Failed", f"SMTP test failed: {e}", parent=parent_window)

    def get_all_external_report_watch_folders(self):
        folders_str = self.config.get("Paths", "external_report_watch_folders_list", fallback="")
        return [os.path.normpath(os.path.expanduser(f.strip())) for f in folders_str.split(';') if f.strip()]

    def _process_auto_send_for_external_report_file(self, file_path):
        # This enhanced method now walks up the directory tree to find the correct settings
        filename = os.path.basename(file_path)
        current_folder = os.path.dirname(file_path)
        
        # Walk up to find a configured parent watch folder
        configured_parent_folder = None
        all_watch_folders = self.get_all_external_report_watch_folders()
        temp_path = current_folder
        while True:
            if temp_path in all_watch_folders:
                configured_parent_folder = temp_path
                break
            parent = os.path.dirname(temp_path)
            if parent == temp_path: # Reached root
                break
            temp_path = parent

        if not configured_parent_folder:
            logging.info(f"No auto-send configuration found for file {file_path} or its parents.")
            return

        folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(configured_parent_folder)}"
        if not self.config.has_section(folder_section) or not self.config.getboolean(folder_section, "enabled", fallback=False):
            return

        recipients = self.config.get(folder_section, "recipients", fallback="")
        if not recipients: return

        template_name = self.config.get(folder_section, "template", fallback="")
        patient_data = {} # extract_patient_data_from_filename(filename) can be used here if needed
        
        if template_name == "Custom":
            subject = self.config.get(folder_section, "custom_subject", fallback="External Report: {Filename}")
            body = self.config.get(folder_section, "custom_body", fallback="Attached: {Filename}")
        else: # Default or named template
            subject = self.config.get("EmailTemplates", f"{template_name}_subject", fallback="Report: {Filename}")
            body = self.config.get("EmailTemplates", f"{template_name}_body", fallback="Attached: {Filename}")
            
        replacements = {
            '{Filename}': filename, '{FolderPath}': current_folder, '{DateTime}': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            '{Patient Name}': patient_data.get('patient_name', 'N/A'), '{Modality}': patient_data.get('modality', 'N/A'),
            '{Study Description}': patient_data.get('description', 'N/A'), '{Date}': patient_data.get('date', 'N/A')
        }
        for placeholder, value in replacements.items():
            subject = subject.replace(placeholder, str(value))
            body = body.replace(placeholder, str(value))
        
        self.send_email_with_report(recipients, subject, body, [file_path], update_recents=False)

    def open_served_worklist_viewer(self):
        ServedWorklistDialog(self.root, self.config, self.current_palette, self.get_ui_label, self)

    # --- NEW AND REFACTORED DICOM WORKFLOW METHODS ---
    def _start_pacs_query_job(self, patient_data):
        if not PYNETDICOM_AVAILABLE:
            activity_logger.warning("Cannot start DICOM workflow: pynetdicom is not available.")
            return

        job_id = patient_data["Accession Number"]
        job_info = {
            "patient_data": patient_data,
            "created_at": datetime.now().isoformat()
        }
        
        # Add job to the persistent queue
        all_jobs = _read_jobs()
        all_jobs[job_id] = job_info
        _write_jobs(all_jobs)
        
        activity_logger.info(f"Job queued for {patient_data['Patient Name']} ({patient_data['Study Description']}). First check in 10 mins.")
        
        # Start the worker thread
        worker_thread = threading.Thread(
            target=self._pacs_query_and_conversion_worker,
            args=(job_id,),
            name=f"DICOM_Worker_{job_id}"
        )
        worker_thread.daemon = True
        worker_thread.start()

    def _resume_pending_jobs(self):
        if not PYNETDICOM_AVAILABLE: return
        
        all_jobs = _read_jobs()
        if not all_jobs: return

        activity_logger.info(f"Resuming {len(all_jobs)} pending DICOM conversion job(s)...")
        for job_id in all_jobs:
            worker_thread = threading.Thread(
                target=self._pacs_query_and_conversion_worker,
                args=(job_id,),
                name=f"DICOM_Worker_{job_id}"
            )
            worker_thread.daemon = True
            worker_thread.start()

    def _pacs_query_and_conversion_worker(self, job_id):
        all_jobs = _read_jobs()
        job_info = all_jobs.get(job_id)
        if not job_info:
            logging.error(f"Worker started for job_id {job_id}, but it was not found in the jobs file.")
            return

        patient_data = job_info["patient_data"]
        accession = patient_data["Accession Number"]
        patient_name = patient_data["Patient Name"]
        study_desc = patient_data["Study Description"]
        modality = patient_data["Modality"]
        pid = patient_data["Patient ID"]

        # Polling parameters
        total_retries = 24  # 4 hours / 10 minutes
        retry_interval_seconds = 600 # 10 minutes

        study_uid = None
        
        # Initial 10-minute delay
        time.sleep(retry_interval_seconds)

        for i in range(total_retries):
            attempt_num = i + 1
            activity_logger.info(f"Checking for {patient_name}'s {study_desc} scan. (Attempt {attempt_num}/{total_retries})")
            
            uid = self._find_study_uid_from_pacs(accession)
            if uid:
                study_uid = uid
                activity_logger.info(f"SUCCESS: Found Study UID for {patient_name}'s scan ({accession}).")
                break
            
            if attempt_num < total_retries:
                activity_logger.info(f"Study for {accession} not found. Next check in 10 minutes.")
                time.sleep(retry_interval_seconds)
        
        try:
            if not study_uid:
                activity_logger.warning(f"TIMEOUT: Could not find study for {patient_name} ({accession}) after 4 hours.")
                return # End of this worker's job

            # --- Proceed with conversion and sending ---
            watch_folder = self.get_modality_specific_path("Paths.WatchFolders.Modalities", modality)
            if not watch_folder or not os.path.isdir(watch_folder):
                activity_logger.error(f"Cannot process files for {accession}: Watch folder for modality '{modality}' is not valid.")
                return

            files_to_process = []
            for dirpath, _, filenames in os.walk(watch_folder):
                for fname in filenames:
                    if pid.lower() in fname.lower() and (fname.lower().endswith(('.pdf', '.jpg', '.jpeg'))):
                        files_to_process.append(os.path.join(dirpath, fname))
            
            if not files_to_process:
                activity_logger.warning(f"Found Study UID for {accession}, but no matching PDF/JPG files were found in the watch folder.")
                return

            activity_logger.info(f"Found {len(files_to_process)} file(s) for {accession}. Starting conversion and sending.")
            success_count = 0
            for fpath in files_to_process:
                ds = self._convert_file_to_dicom(fpath, patient_data, study_uid)
                if ds:
                    if self._send_dicom_file(ds, modality):
                        success_count += 1
                        activity_logger.info(f"Successfully sent: {os.path.basename(fpath)}")
                    else:
                        activity_logger.error(f"Failed to send: {os.path.basename(fpath)}")
                else:
                    activity_logger.error(f"Failed to convert: {os.path.basename(fpath)}")

            activity_logger.info(f"Finished job for {accession}. Sent {success_count}/{len(files_to_process)} files.")

        finally:
            # Clean up: remove the job from the queue file
            all_jobs = _read_jobs()
            if job_id in all_jobs:
                del all_jobs[job_id]
                _write_jobs(all_jobs)
            logging.info(f"Job {job_id} removed from the queue.")
            
    def _find_study_uid_from_pacs(self, accession_number):
        pacs_ae = self.config.get("DICOM.QueryPACS", "ae_title", fallback="")
        pacs_ip = self.config.get("DICOM.QueryPACS", "ip", fallback="")
        pacs_port_str = self.config.get("DICOM.QueryPACS", "port", fallback="")
        if not all([pacs_ae, pacs_ip, pacs_port_str]): return None
        try:
            pacs_port = int(pacs_port_str)
        except ValueError: return None

        ae = AE()
        ae.add_requested_context(StudyRootQueryRetrieveInformationModelFind)
        ds = Dataset()
        ds.QueryRetrieveLevel = 'STUDY'
        ds.AccessionNumber = accession_number
        ds.StudyInstanceUID = ''
        
        try:
            assoc = ae.associate(pacs_ip, pacs_port, ae_title=pacs_ae.encode('ascii'))
            if assoc.is_established:
                responses = assoc.send_c_find(ds, query_model='S')
                for (status, identifier) in responses:
                    if status and status.Status in (0xFF00, 0xFF01) and identifier:
                        assoc.release()
                        return identifier.StudyInstanceUID
                assoc.release()
        except Exception as e:
            logging.error(f"C-FIND exception for Acc {accession_number}: {e}")
        return None

    # ... (_convert_file_to_dicom and _send_dicom_file are unchanged) ...
    def _convert_file_to_dicom(self, file_path, patient_data, study_uid):
        try:
            ds = Dataset()
            ds.PatientName = patient_data.get("Patient Name")
            ds.PatientID = patient_data.get("Patient ID")
            ds.PatientBirthDate = patient_data.get("Date of Birth")
            ds.PatientSex = patient_data.get("Sex")

            ds.StudyInstanceUID = study_uid
            ds.AccessionNumber = patient_data.get("Accession Number")
            ds.StudyDescription = patient_data.get("Study Description")
            ds.StudyDate = patient_data.get("Study Date")
            ds.StudyTime = patient_data.get("Study Time")

            ds.Modality = "OT"
            ds.SOPInstanceUID = pydicom_module.uid.generate_uid()
            ds.SeriesInstanceUID = pydicom_module.uid.generate_uid()

            ds.file_meta = FileMetaDataset()
            ds.file_meta.MediaStorageSOPInstanceUID = ds.SOPInstanceUID
            ds.file_meta.TransferSyntaxUID = pydicom_module.uid.ImplicitVRLittleEndian

            ds.ConversionType = "WSD"
            ds.ContentDate = datetime.now().strftime('%Y%m%d')
            ds.ContentTime = datetime.now().strftime('%H%M%S.%f')

            if file_path.lower().endswith('.pdf'):
                ds.SOPClassUID = EncapsulatedPDFStorage
                ds.file_meta.MediaStorageSOPClassUID = EncapsulatedPDFStorage
                ds.MIMETypeOfEncapsulatedDocument = "application/pdf"
                with open(file_path, 'rb') as f:
                    ds.EncapsulatedDocument = f.read()
            elif file_path.lower().endswith(('.jpg', '.jpeg')):
                ds.SOPClassUID = SecondaryCaptureImageStorage
                ds.file_meta.MediaStorageSOPClassUID = SecondaryCaptureImageStorage
                img = pydicom_module.pixel_data_handlers.pillow_handler.get_image_from_path(file_path)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                ds.Rows, ds.Columns = img.height, img.width
                ds.PhotometricInterpretation = "RGB"
                ds.SamplesPerPixel = 3
                ds.BitsAllocated = 8
                ds.BitsStored = 8
                ds.HighBit = 7
                ds.PixelRepresentation = 0
                ds.PlanarConfiguration = 0
                ds.PixelData = img.tobytes()
            else:
                return None

            pydicom_module.dataset.validate_file_meta(ds.file_meta, enforce_standard=True)
            return ds
        except Exception as e:
            logging.exception(f"Error converting {file_path} to DICOM: {e}")
            return None

    def _send_dicom_file(self, ds, modality):
        dest_section = f"DICOM.Destinations.{modality}"
        dest_ae = self.config.get(dest_section, "ae_title", fallback="")
        dest_ip = self.config.get(dest_section, "ip", fallback="")
        dest_port_str = self.config.get(dest_section, "port", fallback="")
        if not all([dest_ae, dest_ip, dest_port_str]): return False
        try:
            dest_port = int(dest_port_str)
        except ValueError: return False

        ae = AE()
        ae.add_requested_context(ds.SOPClassUID)
        try:
            assoc = ae.associate(dest_ip, dest_port, ae_title=dest_ae.encode('ascii'))
            if assoc.is_established:
                status = assoc.send_c_store(ds)
                assoc.release()
                return status and status.Status == 0x0000
        except Exception as e:
            logging.exception(f"C-STORE exception to {dest_ae}: {e}")
        return False

    def _test_dicom_echo(self, ae_title, ip, port, parent_window):
        # This method is unchanged.
        if not PYNETDICOM_AVAILABLE:
            messagebox.showerror("DICOM Error", "pynetdicom library is not available.", parent=parent_window)
            return
        if not all([ae_title, ip, port]):
            messagebox.showwarning("Input Missing", "AE Title, IP Address, and Port are all required.", parent=parent_window)
            return
        try:
            port_num = int(port)
        except ValueError:
            messagebox.showerror("Invalid Port", f"The port '{port}' is not a valid number.", parent=parent_window)
            return
        logging.info(f"Attempting C-ECHO to {ae_title} at {ip}:{port}")
        ae = AE()
        ae.add_requested_context(Verification)
        try:
            assoc = ae.associate(ip, port_num, ae_title=ae_title.encode('ascii'))
            if assoc.is_established:
                status = assoc.send_c_echo()
                if status and status.Status == 0x0000:
                    messagebox.showinfo("Success", f"DICOM C-ECHO to {ae_title} was successful!", parent=parent_window)
                else:
                    messagebox.showerror("Verification Failed", f"C-ECHO failed. Status: {status.Status if status else 'Unknown'}", parent=parent_window)
                assoc.release()
            else:
                messagebox.showerror("Association Failed", f"Could not associate with {ae_title} at {ip}:{port}.", parent=parent_window)
        except Exception as e:
            messagebox.showerror("Connection Error", f"An error occurred while trying to connect:\n{e}", parent=parent_window)

# --- Dialog Classes (ReportPickerDialog, EmailComposerDialog, ServedWorklistDialog) are unchanged ---
class ReportPickerDialog(tk.Toplevel):
    # This class is unchanged from the previous version.
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        self.title(self.get_ui_label("email_picker_title", "Select Report to Email"))
        self.geometry("900x600")
        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))
        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        ttk.Label(main_frame, text="Select a report to email:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(0,10), anchor=tk.W)
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(expand=True, fill=tk.BOTH, pady=(0,15))
        columns = ("patient_name", "patient_id", "accession_number", "study_description", "modality", "study_date", "docx_path")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        column_widths = {"patient_name": 150, "patient_id": 100, "accession_number": 120, "study_description": 200, "modality": 80, "study_date": 100, "docx_path": 250}
        for col in columns:
            self.tree.heading(col, text=col.replace("_", " ").title(), command=lambda c=col: self.main_app.sort_treeview_column(self.tree, c, False))
            self.tree.column(col, width=column_widths.get(col, 100), minwidth=50)
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        self.tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        self.tree.bind("<Double-1>", self.on_compose_email)
        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10,0))
        self.compose_button = ttk.Button(button_frame, text=self.get_ui_label("email_picker_button", "Compose Email for Selected Report"),
                                        command=self.on_compose_email, state=tk.DISABLED)
        self.compose_button.pack(side=tk.LEFT, padx=(0,10))
        ttk.Button(button_frame, text="Cancel", command=self.destroy).pack(side=tk.RIGHT)
        ttk.Button(button_frame, text="Refresh", command=self.load_reports).pack(side=tk.RIGHT, padx=(0,10))
        self.load_reports()

    def load_reports(self):
        self.tree.delete(*self.tree.get_children())
        try:
            all_data, _ = get_all_patient_records_db("")
            for record in all_data:
                patient_name = record.get("patient_name", "Unknown")
                patient_id = record.get("patient_id", "Unknown")
                modality = record.get("modality", "Unknown")
                modality_output_folder = self.main_app.get_modality_specific_path("Paths.Output.DOCX.Modalities", modality)
                if not modality_output_folder: continue
                safe_pname = "".join(c if c.isalnum() else "_" for c in patient_name)
                safe_pid = patient_id.replace(' ', '_')
                patient_subfolder_name = f"{safe_pname}_{safe_pid}"
                patient_folder = os.path.join(modality_output_folder, patient_subfolder_name)
                if os.path.isdir(patient_folder):
                    for filename in os.listdir(patient_folder):
                        if filename.endswith("_REPORT.docx"):
                            values = (patient_name, patient_id, record.get("accession_number", ""), record.get("study_description", ""), modality, format_date_friendly(record.get("study_date", "")), os.path.join(patient_folder, filename))
                            self.tree.insert("", tk.END, values=values)
        except Exception as e:
            messagebox.showerror("Error Loading Reports", f"Error loading reports: {e}", parent=self)

    def on_tree_select(self, event=None):
        self.compose_button.config(state=tk.NORMAL if self.tree.selection() else tk.DISABLED)

    def on_compose_email(self, event=None):
        selection = self.tree.selection()
        if not selection: return
        values = self.tree.item(selection[0])['values']
        if len(values) < 7 or not os.path.exists(values[6]): return
        report_folder = os.path.dirname(values[6])
        attachment_files = [values[6]] + [os.path.join(report_folder, f) for f in os.listdir(report_folder) if os.path.isfile(os.path.join(report_folder, f)) and f != os.path.basename(values[6])]
        self.destroy()
        EmailComposerDialog(self.master, self.app_config, self.palette, self.get_ui_label, self.main_app,
                          patient_name=values[0], modality=values[4], study_description=values[3],
                          study_date=values[5], attachment_files=attachment_files)

class EmailComposerDialog(tk.Toplevel):
    # This class is unchanged from the previous version.
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref,
                 patient_name="", modality="", study_description="", study_date="", attachment_files=None):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        self.patient_name = patient_name
        self.modality = modality
        self.study_description = study_description
        self.study_date = study_date
        self.attachment_files = attachment_files or []
        self.title(self.get_ui_label("email_composer_title", "Compose Email"))
        self.geometry("800x700")
        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))
        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        main_frame.columnconfigure(1, weight=1)
        recipients_frame = ttk.LabelFrame(main_frame, text=self.get_ui_label("email_composer_select_recipients_label", "Select Recipients:"),
                                        style='Custom.TFrame', padding=10)
        recipients_frame.grid(row=0, column=0, columnspan=2, sticky=tk.EW, pady=(0,10))
        recipients_frame.columnconfigure(1, weight=1)
        ttk.Label(recipients_frame, text=self.get_ui_label("recent_recipients_combobox_label", "Recent/Favorites:")).grid(row=0, column=0, sticky=tk.W, padx=(0,5))
        self.recipients_listbox = tk.Listbox(recipients_frame, selectmode=tk.MULTIPLE, height=4,
                                           bg=self.palette.get("entry_bg"), fg=self.palette.get("entry_fg"),
                                           selectbackground=self.palette.get("header_fg"), selectforeground=self.palette.get("button_fg"))
        self.recipients_listbox.grid(row=0, column=1, sticky=tk.EW, padx=5)
        ttk.Button(recipients_frame, text=self.get_ui_label("email_composer_add_selected_button", "Add Selected to 'To:' Field"),
                  command=self._add_selected_recipients_to_to_field).grid(row=0, column=2, padx=(5,0))
        self.populate_recipient_listbox()
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_to", "To:")).grid(row=1, column=0, sticky=tk.W, pady=(5,0))
        self.to_entry = ttk.Entry(main_frame, width=60)
        self.to_entry.grid(row=1, column=1, sticky=tk.EW, padx=(10,0), pady=(5,0))
        template_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        template_frame.grid(row=2, column=0, columnspan=2, sticky=tk.EW, pady=(10,0))
        template_frame.columnconfigure(1, weight=1)
        ttk.Label(template_frame, text=self.get_ui_label("email_composer_template_label", "Email Template:")).pack(side=tk.LEFT)
        self.template_combo = ttk.Combobox(template_frame, width=30, state="readonly")
        self.template_combo.pack(side=tk.LEFT, padx=(10,0))
        self.template_combo.bind("<<ComboboxSelected>>", self.on_template_selected)
        self.populate_template_combobox()
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_subject", "Subject:")).grid(row=3, column=0, sticky=tk.W, pady=(10,0))
        self.subject_entry = ttk.Entry(main_frame, width=60)
        self.subject_entry.grid(row=3, column=1, sticky=tk.EW, padx=(10,0), pady=(10,0))
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_body", "Body:")).grid(row=4, column=0, sticky=tk.NW, pady=(10,0))
        self.body_text = tk.Text(main_frame, height=15, width=60,
                                bg=self.palette.get("entry_bg"), fg=self.palette.get("entry_fg"),
                                insertbackground=self.palette.get("entry_fg"))
        self.body_text.grid(row=4, column=1, sticky=tk.EW + tk.NS, padx=(10,0), pady=(10,0))
        main_frame.rowconfigure(4, weight=1)
        attachments_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        attachments_frame.grid(row=5, column=0, columnspan=2, sticky=tk.EW, pady=(10,0))
        ttk.Label(attachments_frame, text=self.get_ui_label("email_composer_attachments_label", "Attachments:")).pack(side=tk.LEFT)
        attachment_names = [os.path.basename(f) for f in self.attachment_files]
        ttk.Label(attachments_frame, text=f"{len(attachment_names)} file(s): {', '.join(attachment_names[:3])}{'...' if len(attachment_names) > 3 else ''}",
                 foreground="gray").pack(side=tk.LEFT, padx=(10,0))
        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.grid(row=6, column=0, columnspan=2, sticky=tk.EW, pady=(15,0))
        ttk.Button(button_frame, text=self.get_ui_label("email_composer_send_button", "Send Email"),
                  command=self.send_composed_email).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Cancel", command=self.destroy).pack(side=tk.RIGHT)
        if self.template_combo['values']:
            self.template_combo.set(self.template_combo['values'][0])
            self.apply_email_template()

    def populate_recipient_listbox(self):
        recent_emails = [e.strip() for e in self.app_config.get("EmailRecipients", "recent_list", fallback="").split(';') if e.strip()]
        favorite_emails = [e.strip() for e in self.app_config.get("EmailRecipients", "favorite_list", fallback="").split(';') if e.strip()]
        for email in favorite_emails: self.recipients_listbox.insert(tk.END, f"⭐ {email}")
        for email in recent_emails:
            if email not in favorite_emails: self.recipients_listbox.insert(tk.END, f"📧 {email}")

    def _add_selected_recipients_to_to_field(self):
        selected_emails = [self.recipients_listbox.get(idx).replace("⭐ ", "").replace("📧 ", "") for idx in self.recipients_listbox.curselection()]
        if selected_emails:
            current_to = self.to_entry.get().strip()
            new_to = f"{current_to}; {'; '.join(selected_emails)}" if current_to else '; '.join(selected_emails)
            self.to_entry.delete(0, tk.END)
            self.to_entry.insert(0, new_to)

    def populate_template_combobox(self):
        template_names = [t.strip() for t in self.app_config.get("EmailTemplates", "template_names", fallback="").split(';') if t.strip()]
        if template_names: self.template_combo['values'] = template_names

    def on_template_selected(self, event=None):
        self.apply_email_template()

    def apply_email_template(self):
        template_name = self.template_combo.get()
        if not template_name: return
        subject_template = self.app_config.get("EmailTemplates", f"{template_name}_subject", fallback="")
        body_template = self.app_config.get("EmailTemplates", f"{template_name}_body", fallback="")
        replacements = {
            '{Patient Name}': self.patient_name, '{Modality}': self.modality, '{Study Description}': self.study_description,
            '{Date}': self.study_date, '{Report Filename}': os.path.basename(self.attachment_files[0]) if self.attachment_files else "",
            '{Attachment Count}': str(len(self.attachment_files)), '{All Attachment Names}': ', '.join([os.path.basename(f) for f in self.attachment_files])
        }
        for placeholder, value in replacements.items():
            subject_template = subject_template.replace(placeholder, str(value))
            body_template = body_template.replace(placeholder, str(value))
        self.subject_entry.delete(0, tk.END)
        self.subject_entry.insert(0, subject_template)
        self.body_text.delete("1.0", tk.END)
        self.body_text.insert("1.0", body_template)

    def send_composed_email(self):
        to_emails = self.to_entry.get().strip()
        subject = self.subject_entry.get().strip()
        body = self.body_text.get("1.0", tk.END).strip()
        if not to_emails or not subject:
            messagebox.showerror("Missing Information", "Please enter recipients and a subject.", parent=self)
            return
        if self.main_app.send_email_with_report(to_emails, subject, body, self.attachment_files):
            messagebox.showinfo("Email Sent", "Email sent successfully!", parent=self)
            self.destroy()

class ServedWorklistDialog(tk.Toplevel):
    # This class is unchanged from the previous version.
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        self.original_data = {}
        self.changes_made = False

        self.title(self.get_ui_label("view_served_worklist_title", "Served Worklist Viewer"))
        self.geometry("1400x800")
        self.transient(parent)
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))

        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)

        header_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        header_frame.pack(fill=tk.X, pady=(0,10))

        ttk.Label(header_frame, text="MWL Server Database Entries",
                 font=('Helvetica', 14, 'bold'), style="Header.TLabel").pack(side=tk.LEFT)

        ttk.Button(header_frame, text="Refresh", command=self.refresh_list).pack(side=tk.RIGHT, padx=(10,0))

        self.save_button = ttk.Button(header_frame, text="Save Changes", command=self.save_changes, state=tk.DISABLED)
        self.save_button.pack(side=tk.RIGHT, padx=(10,0))

        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(expand=True, fill=tk.BOTH, pady=(0,10))

        columns = ("id", "patient_name", "patient_id", "accession_number", "dob_yyyymmdd", "sex",
                  "study_date", "study_time", "study_description", "referred_from", "modality",
                  "requesting_physician", "requested_procedure_id", "scheduled_station_ae_title")

        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20, selectmode="extended")

        column_widths = {"id": 50, "patient_name": 150, "patient_id": 100, "accession_number": 120,
                        "dob_yyyymmdd": 90, "sex": 40, "study_date": 80, "study_time": 70,
                        "study_description": 200, "referred_from": 120, "modality": 70,
                        "requesting_physician": 150, "requested_procedure_id": 120, "scheduled_station_ae_title": 100}

        for col in columns:
            self.tree.heading(col, text=col.replace("_", " ").title())
            self.tree.column(col, width=column_widths.get(col, 100), minwidth=50)

        tree_scroll_v = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        tree_scroll_h = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=tree_scroll_v.set, xscrollcommand=tree_scroll_h.set)

        self.tree.grid(row=0, column=0, sticky=tk.NSEW)
        tree_scroll_v.grid(row=0, column=1, sticky=tk.NS)
        tree_scroll_h.grid(row=1, column=0, sticky=tk.EW)

        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)

        self.tree.bind("<Double-1>", self.on_item_double_click)

        self.edit_frame = ttk.LabelFrame(main_frame, text="Edit Record", style='Custom.TFrame', padding=10)
        self.edit_widgets = {}
        self.current_editing_id = None

        self.create_edit_widgets()

        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10,0))

        ttk.Button(button_frame, text="Delete Selected", command=self.delete_selected_records).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Close", command=self.on_closing).pack(side=tk.RIGHT)

        self.refresh_list()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def create_edit_widgets(self):
        fields = {
            "patient_name": (0, 0, 30, False), "patient_id": (0, 2, 20, False),
            "accession_number": (1, 0, 30, False), "dob_yyyymmdd": (1, 2, 20, False),
            "sex": (2, 0, 10, True), "study_date": (2, 2, 20, False),
            "study_time": (3, 0, 20, False), "study_description": (3, 2, 30, False),
            "referred_from": (4, 0, 30, False), "modality": (4, 2, 20, True),
            "requesting_physician": (5, 0, 30, False), "requested_procedure_id": (5, 2, 30, False),
            "scheduled_station_ae_title": (6, 0, 30, False)
        }
        for key, (row, col, width, is_combo) in fields.items():
            label_text = key.replace("_", " ").title() + ":"
            ttk.Label(self.edit_frame, text=label_text).grid(row=row, column=col*2, sticky=tk.W, padx=5, pady=5)
            if is_combo:
                widget = ttk.Combobox(self.edit_frame, width=width, state="readonly" if key == "modality" else "normal")
                if key == "sex": widget['values'] = ["M", "F", "O"]
                widget.bind("<<ComboboxSelected>>", self.on_field_change)
            else:
                widget = ttk.Entry(self.edit_frame, width=width)
                widget.bind("<KeyRelease>", self.on_field_change)
            widget.grid(row=row, column=col*2 + 1, sticky=tk.EW, padx=5, pady=5)
            self.edit_widgets[key] = widget

    def on_field_change(self, event=None):
        self.changes_made = True
        self.save_button.config(state=tk.NORMAL)

    def get_current_edit_data(self):
        data = {}
        for key, widget in self.edit_widgets.items():
            data[key] = widget.get()
        return data

    def on_item_double_click(self, event):
        if self.changes_made and not messagebox.askyesno("Unsaved Changes", "Discard unsaved changes and edit a new record?", parent=self):
            return
        item_id = self.tree.identify_row(event.y)
        if not item_id: return
        self.tree.selection_set(item_id)
        record_id = self.tree.item(item_id, "values")[0]
        self.current_editing_id = record_id
        db_record = get_patient_record_by_db_id(record_id)
        if not db_record:
            messagebox.showerror("Error", f"Could not find record with ID {record_id}.", parent=self)
            return
        self.original_data = db_record
        for key, widget in self.edit_widgets.items():
            value = db_record.get(key, "")
            if isinstance(widget, ttk.Combobox): widget.set(value)
            else:
                widget.delete(0, tk.END)
                widget.insert(0, value)
        self.edit_frame.pack(fill=tk.X, pady=(10,0))
        self.changes_made = False
        self.save_button.config(state=tk.DISABLED)

    def save_changes(self):
        if not self.current_editing_id or not self.changes_made: return
        updated_data = self.get_current_edit_data()
        final_data = self.original_data.copy()
        # Map widget keys to database dictionary keys
        final_data["Patient Name"] = updated_data.get("patient_name")
        final_data["Patient ID"] = updated_data.get("patient_id")
        final_data["Accession Number"] = updated_data.get("accession_number")
        final_data["Date of Birth"] = updated_data.get("dob_yyyymmdd")
        final_data["Sex"] = updated_data.get("sex")
        final_data["Study Date"] = updated_data.get("study_date")
        final_data["Study Time"] = updated_data.get("study_time")
        final_data["Study Description"] = updated_data.get("study_description")
        final_data["Referred From"] = updated_data.get("referred_from")
        final_data["Modality"] = updated_data.get("modality")
        final_data["Requesting Physician"] = updated_data.get("requesting_physician")
        final_data["Requested Procedure ID"] = updated_data.get("requested_procedure_id")
        final_data["Scheduled Station AE Title"] = updated_data.get("scheduled_station_ae_title")

        if update_patient_record_db(self.current_editing_id, final_data):
            messagebox.showinfo("Success", "Record updated successfully.", parent=self)
            self.edit_frame.pack_forget()
            self.current_editing_id = None
            self.changes_made = False
            self.save_button.config(state=tk.DISABLED)
            self.refresh_list()
        else:
            messagebox.showerror("Error", "Failed to update record in the database.", parent=self)

    def refresh_list(self, event=None):
        if self.changes_made and not messagebox.askyesno("Unsaved Changes", "Discard unsaved changes and refresh?", parent=self):
            return
        for i in self.tree.get_children(): self.tree.delete(i)
        all_data, columns = get_all_patient_records_db()
        for record in all_data:
            values = [record.get(col, "") for col in columns]
            self.tree.insert("", tk.END, values=values)
        self.edit_frame.pack_forget()
        self.current_editing_id = None
        self.changes_made = False
        self.save_button.config(state=tk.DISABLED)

    def on_closing(self):
        if self.changes_made and messagebox.askyesno("Unsaved Changes", "Close without saving changes?", parent=self):
            self.destroy()
        elif not self.changes_made:
            self.destroy()

    def delete_selected_records(self):
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("No Selection", "Please select one or more records to delete.", parent=self)
            return
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete {len(selected_items)} record(s)?", parent=self):
            deleted_count = 0
            for item in selected_items:
                record_id = self.tree.item(item, "values")[0]
                if db_execute("DELETE FROM patient_records WHERE id=?", (record_id,), commit=True):
                    deleted_count += 1
            messagebox.showinfo("Success", f"{deleted_count} record(s) deleted.", parent=self)
            self.refresh_list()

# --- Main Execution ---
def main():
    root = tk.Tk()
    try:
        app = PatientRegistrationApp(root)
        root.protocol("WM_DELETE_WINDOW", lambda: on_closing(root, app))
        root.mainloop()
    except Exception as e:
        logging.exception("A critical error occurred during application startup.")
        messagebox.showerror("Fatal Error", f"A critical error occurred: {e}\n\nPlease check the log file for details.")
        if 'root' in locals() and root.winfo_exists():
            root.destroy()

def on_closing(root, app_instance):
    logging.info("on_closing called. Initiating shutdown.")
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        app_instance.shutdown()
        root.destroy()
        logging.info("Application window destroyed. Exiting.")
        sys.exit(0)

if __name__ == "__main__":
    main()
